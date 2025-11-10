"""
Comprehensive DOCX Test Document Generator
Creates a Word document with all possible formatting elements for testing
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set link color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    
    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def create_comprehensive_docx():
    """
    Create a comprehensive DOCX file with all formatting elements
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Comprehensive DOCX Test Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    para = doc.add_paragraph()
    para.add_run('Purpose: ').bold = True
    para.add_run('Test all Word document elements for conversion quality\n')
    para.add_run('Created: ').italic = True
    para.add_run('November 10, 2025 | ')
    para.add_run('Version: ').italic = True
    para.add_run('1.0')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # === Section 1: Headings ===
    doc.add_heading('1. Heading Hierarchy', 1)
    doc.add_heading('Heading Level 1', 1)
    doc.add_heading('Heading Level 2', 2)
    doc.add_heading('Heading Level 3', 3)
    doc.add_heading('Heading Level 4', 4)
    doc.add_heading('Heading Level 5', 5)
    
    # === Section 2: Text Formatting ===
    doc.add_heading('2. Text Formatting', 1)
    
    p = doc.add_paragraph('This paragraph demonstrates ')
    p.add_run('bold text').bold = True
    p.add_run(', ')
    p.add_run('italic text').italic = True
    p.add_run(', ')
    p.add_run('underlined text').underline = True
    p.add_run(', and ')
    run = p.add_run('colored text')
    run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run('.')
    
    p = doc.add_paragraph('This shows ')
    run = p.add_run('highlighted text')
    run.font.highlight_color = 7  # Yellow highlight
    p.add_run(' and ')
    run = p.add_run('strike-through text')
    run.font.strike = True
    p.add_run('.')
    
    p = doc.add_paragraph('Font sizes: ')
    run = p.add_run('Small (8pt) ')
    run.font.size = Pt(8)
    run = p.add_run('Normal (11pt) ')
    run.font.size = Pt(11)
    run = p.add_run('Large (16pt) ')
    run.font.size = Pt(16)
    run = p.add_run('Extra Large (24pt)')
    run.font.size = Pt(24)
    
    p = doc.add_paragraph()
    p.add_run('Superscript: E = mc')
    run = p.add_run('2')
    run.font.superscript = True
    
    p = doc.add_paragraph()
    p.add_run('Subscript: H')
    run = p.add_run('2')
    run.font.subscript = True
    p.add_run('O')
    
    # === Section 3: Lists ===
    doc.add_heading('3. Lists', 1)
    
    doc.add_heading('3.1 Bulleted List', 2)
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    doc.add_paragraph('Third bullet point with sub-items', style='List Bullet')
    doc.add_paragraph('Sub-item 1', style='List Bullet 2')
    doc.add_paragraph('Sub-item 2', style='List Bullet 2')
    doc.add_paragraph('Deep nested item', style='List Bullet 3')
    doc.add_paragraph('Fourth bullet point', style='List Bullet')
    
    doc.add_heading('3.2 Numbered List', 2)
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    doc.add_paragraph('Third numbered item', style='List Number')
    doc.add_paragraph('Sub-numbered item A', style='List Number 2')
    doc.add_paragraph('Sub-numbered item B', style='List Number 2')
    doc.add_paragraph('Fourth numbered item', style='List Number')
    
    # === Section 4: Tables ===
    doc.add_heading('4. Tables', 1)
    
    doc.add_heading('4.1 Simple Table', 2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Age'
    hdr_cells[2].text = 'City'
    hdr_cells[3].text = 'Occupation'
    
    # Data rows
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Alice Johnson'
    row_cells[1].text = '28'
    row_cells[2].text = 'New York'
    row_cells[3].text = 'Software Engineer'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'Bob Smith'
    row_cells[1].text = '34'
    row_cells[2].text = 'London'
    row_cells[3].text = 'Data Scientist'
    
    row_cells = table.rows[3].cells
    row_cells[0].text = 'Charlie Brown'
    row_cells[1].text = '45'
    row_cells[2].text = 'Tokyo'
    row_cells[3].text = 'Product Manager'
    
    doc.add_paragraph()  # Spacing
    
    doc.add_heading('4.2 Styled Table', 2)
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Medium Shading 1 Accent 1'
    
    # Header
    hdr = table2.rows[0].cells
    hdr[0].text = 'Quarter'
    hdr[1].text = 'Revenue'
    hdr[2].text = 'Expenses'
    hdr[3].text = 'Profit'
    
    # Data
    data = [
        ['Q1 2024', '$150,000', '$100,000', '$50,000'],
        ['Q2 2024', '$175,000', '$110,000', '$65,000'],
        ['Q3 2024', '$200,000', '$120,000', '$80,000'],
        ['Q4 2024', '$225,000', '$130,000', '$95,000'],
    ]
    
    for i, row_data in enumerate(data, start=1):
        cells = table2.rows[i].cells
        for j, value in enumerate(row_data):
            cells[j].text = value
    
    # === Section 5: Images and Shapes ===
    doc.add_heading('5. Images', 1)
    doc.add_paragraph('Image placeholder: In a real document, images would be embedded here.')
    # doc.add_picture('image.png', width=Inches(3))  # Uncomment if image exists
    
    doc.add_page_break()
    
    # === Section 6: Code Blocks ===
    doc.add_heading('6. Code Blocks', 1)
    
    doc.add_paragraph('Python code example:')
    code_para = doc.add_paragraph(
        'def fibonacci(n):\n'
        '    if n <= 1:\n'
        '        return n\n'
        '    return fibonacci(n-1) + fibonacci(n-2)\n\n'
        '# Test\n'
        'for i in range(10):\n'
        '    print(f"F({i}) = {fibonacci(i)}")',
        style='No Spacing'
    )
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(9)
    
    # === Section 7: Quotes and Special Paragraphs ===
    doc.add_heading('7. Quotes', 1)
    
    quote = doc.add_paragraph(
        '"The best way to predict the future is to invent it."',
        style='Quote'
    )
    
    para = doc.add_paragraph()
    para.add_run('— Alan Kay').italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # === Section 8: Alignments ===
    doc.add_heading('8. Text Alignment', 1)
    
    p_left = doc.add_paragraph('This paragraph is left-aligned (default).')
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p_center = doc.add_paragraph('This paragraph is center-aligned.')
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_right = doc.add_paragraph('This paragraph is right-aligned.')
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p_justify = doc.add_paragraph(
        'This paragraph is justified. Justified text aligns to both the left and '
        'right margins, adding extra space between words as necessary. This creates '
        'a clean, professional look for longer documents.'
    )
    p_justify.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # === Section 9: Special Characters ===
    doc.add_heading('9. Special Characters', 1)
    
    doc.add_paragraph('Math symbols: sum integral sqrt infinity approx not-equal less-equal greater-equal plus-minus multiply')
    doc.add_paragraph('Currency: $ EUR GBP YEN INR')
    doc.add_paragraph('Arrows: left right up down left-right')
    doc.add_paragraph('Symbols: copyright registered trademark section paragraph bullet circle-bullet triangle-bullet')
    doc.add_paragraph('Greek: alpha beta gamma delta epsilon theta lambda mu pi sigma phi psi omega')
    
    # === Section 10: Page Layout ===
    doc.add_heading('10. Page Breaks and Sections', 1)
    
    doc.add_paragraph('This is the last section before the page break.')
    doc.add_paragraph('The next content will appear on a new page.')
    
    doc.add_page_break()
    
    # === Summary Page ===
    doc.add_heading('Document Summary', 0)
    
    doc.add_paragraph('This comprehensive DOCX document includes:')
    
    features = [
        'All heading levels (1-5)',
        'Text formatting (bold, italic, underline, colors, highlights)',
        'Font sizes and styles',
        'Superscript and subscript',
        'Bulleted and numbered lists (nested)',
        'Tables (simple and styled)',
        'Code blocks with monospace font',
        'Quotes with attribution',
        'Text alignments (left, center, right, justify)',
        'Special characters (math, currency, symbols)',
        'Page breaks',
        'Multiple sections',
    ]
    
    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run('[OK] ')
        run.font.color.rgb = RGBColor(0, 128, 0)
        p.add_run(feature)
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Total: ').bold = True
    para.add_run('10 sections covering 50+ Word document features')
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph('© 2025 ConverterAI | Comprehensive Test Suite v1.0')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save('test_comprehensive.docx')
    print("test_comprehensive.docx created successfully!")
    print(f"  - {len(doc.paragraphs)} paragraphs")
    print(f"  - {len(doc.tables)} tables")
    print(f"  - Multiple formatting styles and elements")

if __name__ == '__main__':
    create_comprehensive_docx()
