"""
HTML converter - handles all HTML conversions
"""
import time
from pathlib import Path
from typing import Optional
from bs4 import BeautifulSoup
from markdownify import markdownify as md_convert
from docx import Document
from docx.shared import Inches, Pt
import markdown2
import re

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

from converters.base import BaseConverter, ConversionResult
from utils.logger import logger


class HTMLConverter(BaseConverter):
    """Convert HTML to other formats"""
    
    def convert(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Route to appropriate conversion method"""
        output_format = Path(output_file).suffix.lower().lstrip('.')
        
        start_time = time.time()
        
        # Validate files
        error = self._validate_files(input_file, output_file)
        if error:
            return self._create_error_result(input_file, error, 'html', output_format)
        
        try:
            if output_format == 'pdf':
                result = self._html_to_pdf(input_file, output_file, **options)
            elif output_format in ['docx', 'doc']:
                result = self._html_to_docx(input_file, output_file, **options)
            elif output_format in ['md', 'markdown']:
                result = self._html_to_markdown(input_file, output_file, **options)
            else:
                return self._create_error_result(
                    input_file,
                    f"Unsupported output format: {output_format}",
                    'html',
                    output_format
                )
            
            processing_time = time.time() - start_time
            result.processing_time = processing_time
            return result
            
        except Exception as e:
            logger.error(f"HTML conversion failed: {e}", exc_info=True)
            return self._create_error_result(input_file, str(e), 'html', output_format)
    
    def _html_to_pdf(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert HTML to PDF"""
        logger.info(f"Converting HTML to PDF: {input_file} -> {output_file}")
        
        try:
            # Try to use WeasyPrint if available (requires GTK on Windows)
            try:
                from weasyprint import HTML as WeasyHTML
                WeasyHTML(filename=input_file).write_pdf(output_file)
                logger.info("Used WeasyPrint for HTML to PDF conversion")
            except (ImportError, OSError) as e:
                # Fallback to simpler conversion using reportlab
                logger.warning(f"WeasyPrint not available ({e}), using reportlab fallback")
                
                # Read and parse HTML
                with open(input_file, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                # Clean problematic HTML elements before parsing
                import re
                
                # CRITICAL: Remove <style> and <script> tags completely (they appear as text in PDF!)
                html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                
                # Remove all problematic attributes from anchor tags
                html_content = re.sub(r'<a[^>]*>', '', html_content)
                html_content = re.sub(r'</a>', '', html_content)
                # Remove id attributes that cause issues
                html_content = re.sub(r'\sid=["\'][^"\']*["\']', '', html_content)
                # Remove name attributes
                html_content = re.sub(r'\sname=["\'][^"\']*["\']', '', html_content)
                # Remove rel attributes (not supported by ReportLab)
                html_content = re.sub(r'\srel=["\'][^"\']*["\']', '', html_content)
                # Remove title attributes from links
                html_content = re.sub(r'\stitle=["\'][^"\']*["\']', '', html_content)
                # Remove class attributes (not supported by ReportLab)
                html_content = re.sub(r'\sclass=["\'][^"\']*["\']', '', html_content)
                # Remove target attributes
                html_content = re.sub(r'\starget=["\'][^"\']*["\']', '', html_content)
                # Remove alt attributes from img tags (can cause issues)
                html_content = re.sub(r'\salt=["\'][^"\']*["\']', '', html_content)
                
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Double-check: Remove any remaining style and script tags from soup
                for tag in soup(['style', 'script', 'meta', 'link']):
                    tag.decompose()
                
                # Create PDF with enhanced formatting
                doc = SimpleDocTemplate(
                    output_file, 
                    pagesize=letter,
                    topMargin=0.75*inch,
                    bottomMargin=0.75*inch,
                    leftMargin=0.75*inch,
                    rightMargin=0.75*inch
                )
                
                # Get and customize styles
                styles = getSampleStyleSheet()
                
                # Enhanced Heading styles
                styles.add(ParagraphStyle(
                    name='CustomH1',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor=colors.HexColor('#1a1a1a'),
                    spaceAfter=16,
                    spaceBefore=12,
                    fontName='Helvetica-Bold',
                    leading=28
                ))
                
                styles.add(ParagraphStyle(
                    name='CustomH2',
                    parent=styles['Heading2'],
                    fontSize=20,
                    textColor=colors.HexColor('#2d2d2d'),
                    spaceAfter=14,
                    spaceBefore=10,
                    fontName='Helvetica-Bold',
                    leading=24
                ))
                
                styles.add(ParagraphStyle(
                    name='CustomH3',
                    parent=styles['Heading3'],
                    fontSize=16,
                    textColor=colors.HexColor('#404040'),
                    spaceAfter=12,
                    spaceBefore=8,
                    fontName='Helvetica-Bold',
                    leading=20
                ))
                
                styles.add(ParagraphStyle(
                    name='CustomH4',
                    parent=styles['Heading4'],
                    fontSize=14,
                    textColor=colors.HexColor('#555555'),
                    spaceAfter=10,
                    spaceBefore=6,
                    fontName='Helvetica-Bold',
                    leading=18
                ))
                
                # Code block style
                styles.add(ParagraphStyle(
                    name='CodeBlock',
                    parent=styles['Code'],
                    fontName='Courier',
                    fontSize=9,
                    textColor=colors.HexColor('#2d2d2d'),
                    backColor=colors.HexColor('#f5f5f5'),
                    borderColor=colors.HexColor('#dddddd'),
                    borderWidth=1,
                    borderPadding=8,
                    leftIndent=20,
                    rightIndent=20,
                    spaceAfter=12,
                    spaceBefore=12,
                    leading=11
                ))
                
                # Blockquote style
                styles.add(ParagraphStyle(
                    name='BlockQuote',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.HexColor('#555555'),
                    leftIndent=30,
                    rightIndent=30,
                    borderColor=colors.HexColor('#0066cc'),
                    borderWidth=3,
                    borderPadding=10,
                    spaceAfter=12,
                    spaceBefore=12,
                    fontName='Helvetica-Oblique'
                ))
                
                # List item style
                styles.add(ParagraphStyle(
                    name='ListItem',
                    parent=styles['Normal'],
                    fontSize=11,
                    leftIndent=25,
                    spaceAfter=6,
                    bulletIndent=10
                ))
                
                # Enhanced body text
                styles.add(ParagraphStyle(
                    name='EnhancedBody',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.HexColor('#333333'),
                    spaceAfter=8,
                    leading=15,
                    alignment=0
                ))
                
                story = []
                
                # CSS/Style content indicators to skip
                css_indicators = [
                    'font-family:',
                    'font-size:',
                    'background:',
                    'background-color:',
                    'padding:',
                    'margin:',
                    'margin-top:',
                    'margin-bottom:',
                    'border:',
                    'border-bottom:',
                    'border-top:',
                    'box-sizing:',
                    'display:',
                    'color:',
                    'line-height:',
                    'max-width:',
                    'text-align:',
                    'font-weight:',
                    '/* ',
                    '*/',
                    '.table',
                    '.text-',
                    '.mb-',
                    '.mt-',
                    '@media',
                    '{\\',  # CSS block start with backslash
                    '}\\',  # CSS block end with backslash
                    'rgba(',
                    ' rem',
                    'px;',
                    '14px',
                    '#fff',
                    '#333',
                    'serif;'
                ]
                
                # Process HTML elements with proper formatting
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'code', 'ul', 'ol', 'blockquote', 'hr', 'table']):
                    try:
                        # Skip CSS/style content that shouldn't be in body
                        element_text = element.get_text().strip()
                        
                        # Skip very short technical content (likely CSS fragments)
                        if len(element_text) < 5 and ('{' in element_text or '}' in element_text or '\\' in element_text):
                            continue
                        
                        # Skip multiline CSS blocks (ending with backslash)
                        if element_text.endswith('\\') and (':' in element_text or '{' in element_text):
                            continue
                        
                        # Skip CSS comments
                        if element_text.startswith('/*') and element_text.endswith('*/'):
                            continue
                        if '/*' in element_text and '*/' in element_text:
                            continue
                        
                        # Skip CSS selector blocks (e.g., "li { font-size: ... }")
                        # Pattern: word/selector { property: value; }
                        import re
                        if re.search(r'[\w\-\.\#\>\:\s,]+\s*\{\s*[\w\-]+\s*:', element_text):
                            continue  # This is a CSS rule
                        
                        # Skip if contains CSS property patterns
                        if re.search(r'[\w\-]+\s*:\s*[\w\-]+\s*;', element_text) and '{' not in element_text:
                            # Check if it's a real CSS property (not "Purpose: Test" or similar)
                            lower = element_text.lower()
                            if any(css_prop in lower for css_prop in ['font-', 'margin', 'padding', 'color:', 'background', 'border', 'width:', 'height:']):
                                continue
                        
                        # Check if this element contains CSS/style code
                        is_css_content = False
                        if element.name == 'p' or element.name == 'ul':
                            lower_text = element_text.lower()
                            # If it contains multiple CSS indicators, likely style content
                            css_matches = sum(1 for indicator in css_indicators if indicator.lower() in lower_text)
                            if css_matches >= 3:  # At least 3 CSS indicators = likely CSS code
                                is_css_content = True
                            # Also check for single strong indicators
                            elif any(ind in lower_text for ind in ['font-family:', 'background-color:', 'box-sizing:', '@media', 'font-feature-settings:', 'border-bottom-color:']):
                                is_css_content = True
                        
                        if is_css_content:
                            continue  # Skip this element
                        
                        # Headings with hierarchy
                        if element.name == 'h1':
                            text = element.get_text().strip()
                            if text:
                                # Clean problematic characters
                                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(text, styles['CustomH1']))
                                story.append(Spacer(1, 0.2 * inch))
                        
                        elif element.name == 'h2':
                            text = element.get_text().strip()
                            if text:
                                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(text, styles['CustomH2']))
                                story.append(Spacer(1, 0.15 * inch))
                        
                        elif element.name == 'h3':
                            text = element.get_text().strip()
                            if text:
                                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(text, styles['CustomH3']))
                                story.append(Spacer(1, 0.12 * inch))
                        
                        elif element.name in ['h4', 'h5', 'h6']:
                            text = element.get_text().strip()
                            if text:
                                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(text, styles['CustomH4']))
                                story.append(Spacer(1, 0.1 * inch))
                        
                        # Code blocks
                        elif element.name == 'pre':
                            code_text = element.get_text().strip()
                            if code_text:
                                code_text = code_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(code_text, styles['CodeBlock']))
                        
                        # Blockquotes
                        elif element.name == 'blockquote':
                            quote_text = element.get_text().strip()
                            if quote_text:
                                # Clean problematic characters
                                quote_text = quote_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(quote_text, styles['BlockQuote']))
                        
                        # Lists
                        elif element.name in ['ul', 'ol']:
                            for li in element.find_all('li', recursive=False):
                                li_text = li.get_text().strip()
                                if li_text:
                                    # Clean problematic characters
                                    li_text = li_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                    if element.name == 'ul':
                                        bullet = '•'
                                    else:
                                        bullet = f"{element.find_all('li', recursive=False).index(li) + 1}."
                                    story.append(Paragraph(f"{bullet} {li_text}", styles['ListItem']))
                        
                        # Horizontal rule
                        elif element.name == 'hr':
                            story.append(Spacer(1, 0.1 * inch))
                            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#dddddd')))
                            story.append(Spacer(1, 0.1 * inch))
                        
                        # Tables
                        elif element.name == 'table':
                            table_data = []
                            headers = element.find_all('th')
                            if headers:
                                # Clean problematic characters in headers
                                clean_headers = []
                                for th in headers:
                                    text = th.get_text().strip()
                                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                    clean_headers.append(text)
                                table_data.append(clean_headers)
                            
                            for tr in element.find_all('tr'):
                                tds = tr.find_all('td')
                                if tds:
                                    # Clean problematic characters in cells
                                    clean_cells = []
                                    for td in tds:
                                        text = td.get_text().strip()
                                        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                        clean_cells.append(text)
                                    table_data.append(clean_cells)
                            
                            if table_data:
                                pdf_table = Table(table_data)
                                pdf_table.setStyle(TableStyle([
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a90e2')),
                                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
                                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                                ]))
                                story.append(pdf_table)
                                story.append(Spacer(1, 0.2 * inch))
                        
                        # Regular paragraphs with inline formatting
                        elif element.name == 'p':
                            para_html = str(element)
                            para_html = para_html.replace('<p>', '').replace('</p>', '')
                            para_html = para_html.replace('<strong>', '<b>').replace('</strong>', '</b>')
                            para_html = para_html.replace('<em>', '<i>').replace('</em>', '</i>')
                            para_html = para_html.replace('<code>', '<font name="Courier" color="#c7254e">').replace('</code>', '</font>')
                            
                            # Strikethrough support (ReportLab uses <strike>)
                            para_html = para_html.replace('<del>', '<strike>').replace('</del>', '</strike>')
                            para_html = para_html.replace('<s>', '<strike>').replace('</s>', '</strike>')
                            
                            # Underline support (ReportLab uses <u>)
                            # Already supported natively by ReportLab
                            
                            # Superscript and subscript support (ReportLab uses <super> and <sub>)
                            para_html = para_html.replace('<sup>', '<super>').replace('</sup>', '</super>')
                            # <sub> is already supported by ReportLab
                            
                            # Remove ALL anchor tags and their problematic attributes
                            import re
                            # Remove all anchor tags completely (they're already removed but check again)
                            para_html = re.sub(r'<a[^>]*>', '', para_html)
                            para_html = re.sub(r'</a>', '', para_html)
                            # Remove img tags (ReportLab has issues with them)
                            para_html = re.sub(r'<img[^>]*/?>', '[image]', para_html)
                            # Remove br tags properly
                            para_html = para_html.replace('<br/>', ' ').replace('<br>', ' ')
                            # Remove any remaining problematic attributes
                            para_html = re.sub(r'\s(id|name|class|rel|title|target)=["\'][^"\']*["\']', '', para_html)
                            
                            text = para_html.strip()
                            if text and text not in ['', ' ']:
                                try:
                                    story.append(Paragraph(text, styles['EnhancedBody']))
                                    story.append(Spacer(1, 0.05 * inch))
                                except Exception as para_err:
                                    # If paragraph fails, add as plain text
                                    logger.warning(f"Paragraph rendering failed, using plain text: {para_err}")
                                    plain_text = element.get_text().strip()
                                    if plain_text:
                                        story.append(Paragraph(plain_text, styles['EnhancedBody']))
                                        story.append(Spacer(1, 0.05 * inch))
                    
                    except Exception as elem_e:
                        logger.warning(f"Error processing HTML element {element.name}: {elem_e}")
                        continue
                
                doc.build(story)
                logger.info("Used ReportLab with enhanced formatting for HTML to PDF conversion")
            
            logger.info(f"Successfully converted HTML to PDF: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'html',
                'pdf'
            )
            
        except Exception as e:
            logger.error(f"HTML to PDF conversion failed: {e}")
            raise
    
    def _html_to_docx(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert HTML to DOCX"""
        logger.info(f"Converting HTML to DOCX: {input_file} -> {output_file}")
        
        try:
            # Read HTML file
            with open(input_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Parse HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove style, script, meta, and link tags (they shouldn't appear in document content)
            for tag in soup(['style', 'script', 'meta', 'link', 'noscript']):
                tag.decompose()
            
            # Create DOCX document
            doc = Document()
            
            # Process body content
            body = soup.find('body')
            if not body:
                body = soup
            
            self._process_html_element(body, doc)
            
            # Save document
            doc.save(output_file)
            
            logger.info(f"Successfully converted HTML to DOCX: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'html',
                'docx'
            )
            
        except Exception as e:
            logger.error(f"HTML to DOCX conversion failed: {e}")
            raise
    
    def _html_to_markdown(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert HTML to Markdown using professional markdownify library"""
        logger.info(f"Converting HTML to Markdown: {input_file} -> {output_file}")
        
        try:
            # Read HTML file with proper encoding
            with open(input_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Parse with BeautifulSoup to clean first
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove style, script, meta, and link tags
            for tag in soup(['style', 'script', 'meta', 'link', 'noscript']):
                tag.decompose()
            
            # Convert cleaned HTML back to string
            html_content = str(soup)
            
            # Use markdownify for professional conversion
            # This is the Python equivalent of Breakdance used by ConvertAI
            markdown_content = md_convert(
                html_content,
                heading_style="ATX",  # Use # for headings
                bullets="-",  # Use - for bullet points
                strong_em_symbol="**",  # Use ** for bold
                strip=['script', 'style'],  # Remove these tags
                escape_asterisks=False,  # Don't escape * in text
                escape_underscores=False,  # Don't escape _ in text
                newline_style="BACKSLASH"  # Use backslash for line breaks
            )
            
            # Clean up excessive whitespace
            lines = markdown_content.split('\n')
            cleaned_lines = []
            prev_empty = False
            
            for line in lines:
                # Remove trailing whitespace
                line = line.rstrip()
                
                # Handle multiple empty lines
                if line == "":
                    if not prev_empty:
                        cleaned_lines.append(line)
                    prev_empty = True
                else:
                    cleaned_lines.append(line)
                    prev_empty = False
            
            # Join and ensure proper spacing
            final_content = '\n'.join(cleaned_lines)
            
            # Write to file with UTF-8-sig encoding
            with open(output_file, 'w', encoding='utf-8-sig') as f:
                f.write(final_content)
            
            logger.info(f"Successfully converted HTML to Markdown: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'html',
                'markdown'
            )
            
        except Exception as e:
            logger.error(f"HTML to Markdown conversion failed: {e}")
            raise
    
    def _process_html_element(self, element, doc, level=0):
        """Recursively process HTML elements and add to DOCX with enhanced formatting"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for child in element.children:
            if child.name is None:
                # Text node - skip if only whitespace
                text = str(child).strip()
                if text and not text.isspace():
                    # Only add standalone text if not in a structural element
                    if level == 0:
                        doc.add_paragraph(text)
            
            elif child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Heading
                level_num = int(child.name[1])
                text = child.get_text().strip()
                if text:
                    para = doc.add_paragraph(text)
                    if level_num <= 3:
                        para.style = f'Heading {level_num}'
                    elif level_num == 4:
                        para.style = 'Heading 3'
                        # Make it slightly smaller
                        for run in para.runs:
                            run.font.size = Pt(12)
                    else:
                        # H5, H6 - use Heading 3 style but smaller
                        para.style = 'Heading 3'
                        for run in para.runs:
                            run.font.size = Pt(11)
            
            elif child.name == 'p':
                # Paragraph with inline formatting support
                para = doc.add_paragraph()
                self._add_formatted_text(child, para)
            
            elif child.name in ['ul', 'ol']:
                # List with proper bullet/number style
                is_numbered = child.name == 'ol'
                for li in child.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        style = 'List Number' if is_numbered else 'List Bullet'
                        para = doc.add_paragraph(text, style=style)
            
            elif child.name == 'blockquote':
                # Blockquote - italic with indent
                text = child.get_text().strip()
                if text:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.left_indent = Pt(36)
                    para.paragraph_format.right_indent = Pt(36)
                    for run in para.runs:
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(85, 85, 85)
            
            elif child.name == 'table':
                # Table with proper structure
                rows = child.find_all('tr')
                if rows:
                    # Count max columns
                    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                    
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < max_cols:
                                cell_text = cell.get_text().strip()
                                table_cell = table.rows[i].cells[j]
                                table_cell.text = cell_text
                                
                                # Bold and color header row
                                if cell.name == 'th' or i == 0:
                                    for paragraph in table_cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(0, 102, 204)
            
            elif child.name in ['pre', 'code']:
                # Code block
                text = child.get_text()
                if text.strip():
                    para = doc.add_paragraph(text)
                    para.style = 'No Spacing'
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(45, 45, 45)
            
            elif child.name == 'hr':
                # Horizontal rule - add paragraph with bottom border
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.space_before = Pt(6)
            
            elif child.name == 'strong' or child.name == 'b':
                # Bold text in isolation - should be handled by parent
                text = child.get_text().strip()
                if text and level == 0:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.bold = True
            
            elif child.name == 'em' or child.name == 'i':
                # Italic text in isolation - should be handled by parent
                text = child.get_text().strip()
                if text and level == 0:
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.italic = True
            
            elif child.name in ['div', 'section', 'article', 'main', 'aside', 'header', 'footer', 'nav']:
                # Container elements - recurse
                self._process_html_element(child, doc, level + 1)
            
            elif child.name == 'br':
                # Line break - add empty paragraph
                if level == 0:
                    doc.add_paragraph()
    
    def _add_formatted_text(self, element, para):
        """Add text with inline formatting (bold, italic, code, underline, strikethrough, super/subscript, color) to paragraph"""
        from docx.shared import Pt, RGBColor
        import re
        
        for item in element.children:
            if item.name is None:
                # Plain text
                text = str(item)
                if text:
                    para.add_run(text)
            
            elif item.name in ['strong', 'b']:
                # Bold text
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.bold = True
            
            elif item.name in ['em', 'i']:
                # Italic text
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.italic = True
            
            elif item.name == 'u':
                # Underline text
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.font.underline = True
            
            elif item.name in ['del', 's', 'strike']:
                # Strikethrough text
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.font.strike = True
            
            elif item.name == 'sup':
                # Superscript text
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.font.superscript = True
            
            elif item.name == 'sub':
                # Subscript text
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.font.subscript = True
            
            elif item.name == 'code':
                # Inline code
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(199, 37, 78)  # Pink-red
            
            elif item.name == 'a':
                # Link - add as blue underlined text
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    run.font.color.rgb = RGBColor(0, 102, 204)
                    run.font.underline = True
            
            elif item.name == 'span':
                # Span - check for inline styles (color, background)
                text = item.get_text()
                if text:
                    run = para.add_run(text)
                    
                    # Parse inline styles
                    style = item.get('style', '')
                    if style:
                        # Text color
                        color_match = re.search(r'color:\s*([^;]+)', style)
                        if color_match:
                            color_value = color_match.group(1).strip()
                            # Parse hex color (#RRGGBB)
                            if color_value.startswith('#'):
                                hex_color = color_value.lstrip('#')
                                if len(hex_color) == 6:
                                    try:
                                        r = int(hex_color[0:2], 16)
                                        g = int(hex_color[2:4], 16)
                                        b = int(hex_color[4:6], 16)
                                        run.font.color.rgb = RGBColor(r, g, b)
                                    except ValueError:
                                        pass
                            # Parse rgb(r, g, b)
                            elif color_value.startswith('rgb('):
                                rgb_match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_value)
                                if rgb_match:
                                    r, g, b = int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3))
                                    run.font.color.rgb = RGBColor(r, g, b)
                        
                        # Background highlight
                        bg_match = re.search(r'background-color:\s*([^;]+)', style)
                        if bg_match:
                            bg_value = bg_match.group(1).strip()
                            # Map common colors to WD_COLOR_INDEX
                            if bg_value in ['yellow', '#ffff00', '#ff0']:
                                run.font.highlight_color = 7  # Yellow
                            elif bg_value in ['cyan', 'aqua', '#00ffff', '#0ff']:
                                run.font.highlight_color = 11  # Turquoise
                            elif bg_value in ['lime', '#00ff00', '#0f0']:
                                run.font.highlight_color = 6  # Bright green
                
                # Check for nested formatting
                if item.find():
                    self._add_formatted_text(item, para)
            
            else:
                # Unknown element - get text
                text = item.get_text()
                if text:
                    para.add_run(text)
    
    def _html_to_markdown_recursive(self, element, lines, level=0):
        """Recursively convert HTML elements to Markdown with improved formatting"""
        for child in element.children:
            if child.name is None:
                # Text node - skip if only whitespace
                text = str(child).strip()
                if text and not text.isspace():
                    # Only add if not already in a structural element
                    if len(lines) == 0 or (lines and lines[-1] != ""):
                        lines.append(text)
            
            elif child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Heading - add proper spacing
                if lines and lines[-1] != "":
                    lines.append("")
                level_num = int(child.name[1])
                text = child.get_text().strip()
                if text:
                    lines.append(f"{'#' * level_num} {text}")
                    lines.append("")
            
            elif child.name == 'p':
                # Paragraph
                text = child.get_text().strip()
                if text:
                    lines.append(text)
                    lines.append("")
            
            elif child.name in ['strong', 'b']:
                # Bold text
                text = child.get_text().strip()
                if text:
                    lines.append(f"**{text}**")
            
            elif child.name in ['em', 'i']:
                # Italic text
                text = child.get_text().strip()
                if text:
                    lines.append(f"*{text}*")
            
            elif child.name == 'u':
                # Underline text (HTML tag)
                text = child.get_text().strip()
                if text:
                    lines.append(f"<u>{text}</u>")
            
            elif child.name in ['del', 's', 'strike']:
                # Strikethrough text
                text = child.get_text().strip()
                if text:
                    lines.append(f"~~{text}~~")
            
            elif child.name == 'sup':
                # Superscript text
                text = child.get_text().strip()
                if text:
                    lines.append(f"<sup>{text}</sup>")
            
            elif child.name == 'sub':
                # Subscript text
                text = child.get_text().strip()
                if text:
                    lines.append(f"<sub>{text}</sub>")
            
            elif child.name == 'ul':
                # Unordered list - add spacing before
                if lines and lines[-1] != "":
                    lines.append("")
                for li in child.find_all('li', recursive=False):
                    text = ' '.join(li.get_text().split())  # Clean whitespace
                    if text:
                        lines.append(f"- {text}")
                lines.append("")
            
            elif child.name == 'ol':
                # Ordered list - add spacing before
                if lines and lines[-1] != "":
                    lines.append("")
                for idx, li in enumerate(child.find_all('li', recursive=False), 1):
                    text = ' '.join(li.get_text().split())  # Clean whitespace
                    if text:
                        lines.append(f"{idx}. {text}")
                lines.append("")
            
            elif child.name == 'blockquote':
                # Blockquote
                if lines and lines[-1] != "":
                    lines.append("")
                text = child.get_text().strip()
                if text:
                    for line in text.split('\n'):
                        if line.strip():
                            lines.append(f"> {line.strip()}")
                lines.append("")
            
            elif child.name == 'table':
                # Table - improved formatting
                if lines and lines[-1] != "":
                    lines.append("")
                
                rows = child.find_all('tr')
                if rows:
                    for row_idx, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        # Clean cell text - remove extra whitespace
                        cell_texts = [' '.join(cell.get_text().split()) for cell in cells]
                        lines.append("| " + " | ".join(cell_texts) + " |")
                        
                        # Add separator after header
                        if row_idx == 0:
                            separator = "| " + " | ".join(["---"] * len(cell_texts)) + " |"
                            lines.append(separator)
                lines.append("")
            
            elif child.name in ['pre', 'code']:
                # Code block
                if lines and lines[-1] != "":
                    lines.append("")
                code_text = child.get_text()
                if child.name == 'pre':
                    lines.append("```")
                    lines.append(code_text.rstrip())
                    lines.append("```")
                else:
                    lines.append(f"`{code_text}`")
                lines.append("")
            
            elif child.name == 'a':
                # Link
                text = child.get_text().strip()
                href = child.get('href', '')
                if text and href:
                    lines.append(f"[{text}]({href})")
            
            elif child.name == 'img':
                # Image
                alt = child.get('alt', '')
                src = child.get('src', '')
                if src:
                    if lines and lines[-1] != "":
                        lines.append("")
                    lines.append(f"![{alt}]({src})")
                    lines.append("")
            
            elif child.name == 'hr':
                # Horizontal rule
                if lines and lines[-1] != "":
                    lines.append("")
                lines.append("---")
                lines.append("")
            
            elif child.name == 'br':
                # Line break
                lines.append("")
            
            elif child.name in ['div', 'section', 'article', 'main', 'body', 'header', 'footer', 'nav', 'aside']:
                # Container elements - recurse without adding extra lines
                self._html_to_markdown_recursive(child, lines, level + 1)
