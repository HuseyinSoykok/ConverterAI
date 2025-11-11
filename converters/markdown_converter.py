"""
Markdown converter - handles all Markdown conversions
"""
import time
import os
from pathlib import Path
from typing import Optional
import markdown2
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import re

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

from converters.base import BaseConverter, ConversionResult
from utils.logger import logger


class MarkdownConverter(BaseConverter):
    """Convert Markdown to other formats"""
    
    def convert(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Route to appropriate conversion method"""
        output_format = Path(output_file).suffix.lower().lstrip('.')
        
        start_time = time.time()
        
        # Validate files
        error = self._validate_files(input_file, output_file)
        if error:
            return self._create_error_result(input_file, error, 'markdown', output_format)
        
        try:
            if output_format == 'pdf':
                result = self._markdown_to_pdf(input_file, output_file, **options)
            elif output_format in ['docx', 'doc']:
                result = self._markdown_to_docx(input_file, output_file, **options)
            elif output_format in ['html', 'htm']:
                result = self._markdown_to_html(input_file, output_file, **options)
            else:
                return self._create_error_result(
                    input_file,
                    f"Unsupported output format: {output_format}",
                    'markdown',
                    output_format
                )
            
            processing_time = time.time() - start_time
            result.processing_time = processing_time
            return result
            
        except Exception as e:
            logger.error(f"Markdown conversion failed: {e}", exc_info=True)
            return self._create_error_result(input_file, str(e), 'markdown', output_format)
    
    def _markdown_to_html(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert Markdown to HTML using markdown2 with professional ConvertAI-inspired styling"""
        logger.info(f"Converting Markdown to HTML: {input_file} -> {output_file}")
        
        try:
            # Read markdown file with proper encoding
            with open(input_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert to HTML with markdown2 (ConvertAI's markdown-it equivalent)
            # Enable typography like ConvertAI: smart quotes, em/en dashes, etc.
            html_body = markdown2.markdown(
                md_content,
                extras=[
                    "tables",                  # GitHub-style tables
                    "fenced-code-blocks",      # ```code``` blocks
                    "code-friendly",           # Better code handling
                    "cuddled-lists",           # Lists without blank lines (like ConvertAI)
                    "footnotes",               # [^1] footnotes
                    "header-ids",              # Add IDs to headers (like markdown-it-toc)
                    "strike",                  # ~~strikethrough~~
                    "task_list",               # - [ ] task lists (like markdown-it-checkbox)
                    "break-on-newline",        # Line breaks (ConvertAI's breaks:true)
                    "target-blank-links",      # Open external links in new tab
                    "toc",                     # Table of contents support
                    "spoiler",                 # ||spoiler text||
                    "smarty-pants"             # Smart typography: quotes, dashes (ConvertAI's typographer:true)
                ]
            )
            
            # Post-processing: Clean up whitespace (ConvertAI-inspired)
            # Remove excessive newlines, normalize spacing
            import re
            html_body = re.sub(r'\n\n\n+', '\n\n', html_body)  # Max 2 newlines
            html_body = re.sub(r'<p>\s*</p>', '', html_body)  # Remove empty paragraphs
            html_body = re.sub(r'<p>(\s+)', '<p>', html_body)  # Trim paragraph starts
            html_body = re.sub(r'(\s+)</p>', '</p>', html_body)  # Trim paragraph ends
            
            # Load professional CSS separately (avoids f-string brace issues)
            css_path = os.path.join(os.path.dirname(__file__), '..', 'static', 'css', 'export.css')
            try:
                with open(css_path, 'r', encoding='utf-8') as css_file:
                    custom_css = css_file.read()
            except FileNotFoundError:
                logger.warning(f"CSS file not found: {css_path}, using minimal fallback CSS")
                # Minimal fallback CSS if export.css not found
                custom_css = """
body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif;
    line-height: 1.6;
    max-width: 900px;
    margin: 40px auto;
    padding: 20px;
    color: #24292e;
}
table {
    border-collapse: collapse;
    width: 100%;
    margin: 20px 0;
}
th, td {
    border: 1px solid #ddd;
    padding: 12px;
    text-align: left;
}
th {
    background-color: #f6f8fa;
    font-weight: 600;
}
code {
    background: #f6f8fa;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: 'Courier New', monospace;
}
pre {
    background: #f6f8fa;
    padding: 16px;
    border-radius: 6px;
    overflow-x: auto;
}
"""
            
            # Create full HTML document with professional styling  
            # CSS is loaded as a string variable, so no f-string brace conflicts
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Markdown</title>
    <style>
{custom_css}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""
            
            # Write to file with UTF-8-sig encoding
            with open(output_file, 'w', encoding='utf-8-sig') as f:
                f.write(html_content)
            
            logger.info(f"Successfully converted Markdown to HTML: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'markdown',
                'html'
            )
            
        except Exception as e:
            logger.error(f"Markdown to HTML conversion failed: {e}")
            raise
    
    def _markdown_to_pdf(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert Markdown to PDF"""
        logger.info(f"Converting Markdown to PDF: {input_file} -> {output_file}")
        
        try:
            # First convert to HTML
            temp_html = input_file + '.temp.html'
            result = self._markdown_to_html(input_file, temp_html, **options)
            
            if not result.success:
                return result
            
            # Convert HTML to PDF
            try:
                from weasyprint import HTML
                HTML(temp_html).write_pdf(output_file)
                logger.info("Used WeasyPrint for Markdown to PDF conversion")
            except (ImportError, OSError) as e:
                # Fallback to reportlab
                logger.warning(f"WeasyPrint not available ({e}), using reportlab fallback")
                import re  # Re-import in this scope for safety
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                
                # Register Unicode-compatible fonts for Turkish characters
                try:
                    # Try to use DejaVu fonts (widely available and Unicode-complete)
                    # These fonts support Turkish characters: ı, İ, ş, Ş, ğ, Ğ, ö, Ö, ü, Ü, ç, Ç
                    pdfmetrics.registerFont(TTFont('DejaVu', 'DejaVuSans.ttf'))
                    pdfmetrics.registerFont(TTFont('DejaVu-Bold', 'DejaVuSans-Bold.ttf'))
                    pdfmetrics.registerFont(TTFont('DejaVu-Italic', 'DejaVuSans-Oblique.ttf'))
                    pdfmetrics.registerFont(TTFont('DejaVu-BoldItalic', 'DejaVuSans-BoldOblique.ttf'))
                    default_font = 'DejaVu'
                    default_font_bold = 'DejaVu-Bold'
                    logger.info("Using DejaVu fonts for Unicode support")
                except:
                    # Fallback: Try Arial (available on Windows)
                    try:
                        pdfmetrics.registerFont(TTFont('Arial-Unicode', 'arial.ttf'))
                        pdfmetrics.registerFont(TTFont('Arial-Unicode-Bold', 'arialbd.ttf'))
                        default_font = 'Arial-Unicode'
                        default_font_bold = 'Arial-Unicode-Bold'
                        logger.info("Using Arial fonts for Unicode support")
                    except:
                        # Last resort: Use Helvetica but warn about potential issues
                        default_font = 'Helvetica'
                        default_font_bold = 'Helvetica-Bold'
                        logger.warning("Unicode fonts not available - Turkish characters may not display correctly")
                
                # Read HTML and extract text
                with open(temp_html, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Create PDF with proper formatting
                doc = SimpleDocTemplate(
                    output_file, 
                    pagesize=letter,
                    topMargin=0.75*inch,
                    bottomMargin=0.75*inch,
                    leftMargin=0.75*inch,
                    rightMargin=0.75*inch
                )
                
                # Get and customize styles
                styles = getSampleStyleSheet()
                
                # Enhanced Heading 1 - Large, bold, dark blue
                styles.add(ParagraphStyle(
                    name='CustomHeading1',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor=colors.HexColor('#1a1a1a'),
                    spaceAfter=16,
                    spaceBefore=12,
                    fontName=default_font_bold,
                    leading=28
                ))
                
                # Enhanced Heading 2 - Medium, bold, dark gray
                styles.add(ParagraphStyle(
                    name='CustomHeading2',
                    parent=styles['Heading2'],
                    fontSize=20,
                    textColor=colors.HexColor('#2d2d2d'),
                    spaceAfter=14,
                    spaceBefore=10,
                    fontName=default_font_bold,
                    leading=24
                ))
                
                # Enhanced Heading 3 - Smaller, bold
                styles.add(ParagraphStyle(
                    name='CustomHeading3',
                    parent=styles['Heading3'],
                    fontSize=16,
                    textColor=colors.HexColor('#404040'),
                    spaceAfter=12,
                    spaceBefore=8,
                    fontName=default_font_bold,
                    leading=20
                ))
                
                # Enhanced Heading 4 - Smaller
                styles.add(ParagraphStyle(
                    name='CustomHeading4',
                    parent=styles['Heading4'],
                    fontSize=14,
                    textColor=colors.HexColor('#555555'),
                    spaceAfter=10,
                    spaceBefore=6,
                    fontName=default_font_bold,
                    leading=18
                ))
                
                # Enhanced Heading 5 & 6
                styles.add(ParagraphStyle(
                    name='CustomHeading5',
                    parent=styles['Normal'],
                    fontSize=12,
                    textColor=colors.HexColor('#666666'),
                    spaceAfter=8,
                    spaceBefore=6,
                    fontName=default_font_bold,
                    leading=16
                ))
                
                # Code block style
                styles.add(ParagraphStyle(
                    name='CodeBlock',
                    parent=styles['Code'],
                    fontName='Courier',
                    fontSize=9,
                    textColor=colors.HexColor('#2d2d2d'),
                    backColor=colors.HexColor('#f5f5f5'),
                    borderColor=colors.HexColor('#dddddd'),
                    borderWidth=1,
                    borderPadding=8,
                    leftIndent=20,
                    rightIndent=20,
                    spaceAfter=12,
                    spaceBefore=12,
                    leading=11
                ))
                
                # Inline code style
                styles.add(ParagraphStyle(
                    name='InlineCode',
                    parent=styles['Normal'],
                    fontName='Courier',
                    fontSize=10,
                    textColor=colors.HexColor('#c7254e'),
                    backColor=colors.HexColor('#f9f2f4')
                ))
                
                # Blockquote style
                styles.add(ParagraphStyle(
                    name='BlockQuote',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.HexColor('#555555'),
                    leftIndent=30,
                    rightIndent=30,
                    borderColor=colors.HexColor('#0066cc'),
                    borderWidth=3,
                    borderPadding=10,
                    spaceAfter=12,
                    spaceBefore=12,
                    fontName=default_font
                ))
                
                # List item style
                styles.add(ParagraphStyle(
                    name='ListItem',
                    parent=styles['Normal'],
                    fontSize=11,
                    fontName=default_font,
                    leftIndent=25,
                    spaceAfter=6,
                    bulletIndent=10
                ))
                
                # Enhanced body text
                styles.add(ParagraphStyle(
                    name='EnhancedBody',
                    parent=styles['Normal'],
                    fontSize=11,
                    fontName=default_font,
                    textColor=colors.HexColor('#333333'),
                    spaceAfter=8,
                    leading=15,
                    alignment=0  # Left aligned
                ))
                
                story = []
                
                # Process inline formatting: strikethrough, underline, superscript, subscript
                # Convert markdown ~~text~~ to HTML <strike>
                html_content = str(soup)
                html_content = re.sub(r'~~([^~]+)~~', r'<strike>\1</strike>', html_content)
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Process HTML elements properly
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'code', 'ul', 'ol', 'blockquote', 'hr', 'table']):
                    try:
                        # Headings with proper hierarchy
                        if element.name == 'h1':
                            text = element.get_text().strip()
                            if text:
                                story.append(Paragraph(text, styles['CustomHeading1']))
                                story.append(Spacer(1, 0.2 * inch))
                        
                        elif element.name == 'h2':
                            text = element.get_text().strip()
                            if text:
                                story.append(Paragraph(text, styles['CustomHeading2']))
                                story.append(Spacer(1, 0.15 * inch))
                        
                        elif element.name == 'h3':
                            text = element.get_text().strip()
                            if text:
                                story.append(Paragraph(text, styles['CustomHeading3']))
                                story.append(Spacer(1, 0.12 * inch))
                        
                        elif element.name == 'h4':
                            text = element.get_text().strip()
                            if text:
                                story.append(Paragraph(text, styles['CustomHeading4']))
                                story.append(Spacer(1, 0.1 * inch))
                        
                        elif element.name in ['h5', 'h6']:
                            text = element.get_text().strip()
                            if text:
                                story.append(Paragraph(text, styles['CustomHeading5']))
                                story.append(Spacer(1, 0.08 * inch))
                        
                        # Code blocks
                        elif element.name == 'pre':
                            code_text = element.get_text().strip()
                            if code_text:
                                # Escape XML/HTML characters for ReportLab
                                code_text = code_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(code_text, styles['CodeBlock']))
                        
                        # Blockquotes
                        elif element.name == 'blockquote':
                            quote_text = element.get_text().strip()
                            if quote_text:
                                story.append(Paragraph(quote_text, styles['BlockQuote']))
                        
                        # Lists (bullet and numbered)
                        elif element.name in ['ul', 'ol']:
                            for li in element.find_all('li', recursive=False):
                                li_text = li.get_text().strip()
                                if li_text:
                                    if element.name == 'ul':
                                        bullet = '•'
                                    else:
                                        bullet = f"{element.find_all('li', recursive=False).index(li) + 1}."
                                    
                                    story.append(Paragraph(f"{bullet} {li_text}", styles['ListItem']))
                        
                        # Horizontal rule
                        elif element.name == 'hr':
                            story.append(Spacer(1, 0.1 * inch))
                            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#dddddd')))
                            story.append(Spacer(1, 0.1 * inch))
                        
                        # Tables
                        elif element.name == 'table':
                            table_data = []
                            # Get headers
                            headers = element.find_all('th')
                            if headers:
                                table_data.append([th.get_text().strip() for th in headers])
                            
                            # Get rows
                            for tr in element.find_all('tr'):
                                tds = tr.find_all('td')
                                if tds:
                                    table_data.append([td.get_text().strip() for td in tds])
                            
                            if table_data:
                                # Create ReportLab table
                                pdf_table = Table(table_data)
                                pdf_table.setStyle(TableStyle([
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a90e2')),
                                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                    ('FONTNAME', (0, 0), (-1, 0), default_font_bold),
                                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                                    ('FONTNAME', (0, 1), (-1, -1), default_font),
                                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
                                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                                ]))
                                story.append(pdf_table)
                                story.append(Spacer(1, 0.2 * inch))
                        
                        # Regular paragraphs with inline formatting
                        elif element.name == 'p':
                            # Get text with basic HTML formatting preserved
                            para_html = str(element)
                            
                            # Convert bold, italic, code, links, strikethrough, underline, superscript, subscript
                            para_html = para_html.replace('<p>', '').replace('</p>', '')
                            para_html = para_html.replace('<strong>', '<b>').replace('</strong>', '</b>')
                            para_html = para_html.replace('<em>', '<i>').replace('</em>', '</i>')
                            para_html = para_html.replace('<code>', f'<font name="{default_font}" color="#c7254e" backColor="#f9f2f4">').replace('</code>', '</font>')
                            
                            # Strikethrough support (ReportLab uses <strike>)
                            para_html = para_html.replace('<del>', '<strike>').replace('</del>', '</strike>')
                            para_html = para_html.replace('<s>', '<strike>').replace('</s>', '</strike>')
                            
                            # Underline support (ReportLab uses <u>)
                            # Already supported by ReportLab
                            
                            # Superscript and subscript support (ReportLab uses <super> and <sub>)
                            para_html = para_html.replace('<sup>', '<super>').replace('</sup>', '</super>')
                            # <sub> is already supported by ReportLab
                            
                            # Handle links
                            import re
                            para_html = re.sub(r'<a href="([^"]+)">([^<]+)</a>', r'<font color="blue"><u>\2</u></font> (\1)', para_html)
                            
                            text = para_html.strip()
                            if text and text not in ['', ' ']:
                                story.append(Paragraph(text, styles['EnhancedBody']))
                                story.append(Spacer(1, 0.05 * inch))
                    
                    except Exception as e:
                        logger.warning(f"Error processing element {element.name}: {e}")
                        continue
                
                # Build PDF
                doc.build(story)
                logger.info("Used ReportLab with enhanced formatting for Markdown to PDF conversion")
            
            # Clean up temp file
            Path(temp_html).unlink(missing_ok=True)
            
            logger.info(f"Successfully converted Markdown to PDF: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'markdown',
                'pdf'
            )
            
        except Exception as e:
            logger.error(f"Markdown to PDF conversion failed: {e}")
            # Clean up temp file on error
            Path(temp_html).unlink(missing_ok=True)
            raise
    
    def _markdown_to_docx(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert Markdown to DOCX"""
        logger.info(f"Converting Markdown to DOCX: {input_file} -> {output_file}")
        
        try:
            # Read markdown file
            with open(input_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Create DOCX document
            doc = Document()
            
            # Parse markdown line by line
            lines = md_content.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i]
                
                # Heading
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    
                    if text:
                        para = doc.add_paragraph(text)
                        if level == 1:
                            para.style = 'Heading 1'
                        elif level == 2:
                            para.style = 'Heading 2'
                        elif level == 3:
                            para.style = 'Heading 3'
                        else:
                            para.style = 'Heading 4'
                
                # Code block
                elif line.startswith('```'):
                    code_lines = []
                    i += 1
                    while i < len(lines) and not lines[i].startswith('```'):
                        code_lines.append(lines[i])
                        i += 1
                    
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        para = doc.add_paragraph(code_text)
                        para.style = 'No Spacing'
                        # Set monospace font
                        for run in para.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                
                # Bullet list
                elif line.strip().startswith(('- ', '* ', '+ ')):
                    text = line.strip()[2:].strip()
                    doc.add_paragraph(text, style='List Bullet')
                
                # Numbered list
                elif re.match(r'^\d+\.\s', line.strip()):
                    text = re.sub(r'^\d+\.\s', '', line.strip())
                    doc.add_paragraph(text, style='List Number')
                
                # Regular paragraph
                elif line.strip():
                    para = doc.add_paragraph()
                    self._add_markdown_inline_formatting(line.strip(), para)
                
                i += 1
            
            # Save document
            doc.save(output_file)
            
            logger.info(f"Successfully converted Markdown to DOCX: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'markdown',
                'docx'
            )
            
        except Exception as e:
            logger.error(f"Markdown to DOCX conversion failed: {e}")
            raise
    
    def _add_markdown_inline_formatting(self, text: str, para):
        """Parse and add markdown inline formatting to paragraph"""
        from docx.shared import Pt, RGBColor
        from bs4 import BeautifulSoup
        
        # Convert markdown to HTML for easier parsing
        # Strikethrough: ~~text~~
        text = re.sub(r'~~([^~]+)~~', r'<del>\1</del>', text)
        
        # Bold and italic: ***text*** or **text** or *text*
        text = re.sub(r'\*\*\*([^\*]+)\*\*\*', r'<strong><em>\1</em></strong>', text)
        text = re.sub(r'\*\*([^\*]+)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'\*([^\*]+)\*', r'<em>\1</em>', text)
        
        # Inline code: `code`
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        
        # Already HTML tags: <u>, <sup>, <sub>, <span>
        # Wrap in a root element for BeautifulSoup
        html = f'<p>{text}</p>'
        soup = BeautifulSoup(html, 'html.parser')
        
        # Process each element
        for element in soup.find('p').children:
            if element.name is None:
                # Plain text
                para.add_run(str(element))
            
            elif element.name == 'strong':
                # Check if has nested em
                if element.find('em'):
                    run = para.add_run(element.get_text())
                    run.bold = True
                    run.italic = True
                else:
                    run = para.add_run(element.get_text())
                    run.bold = True
            
            elif element.name == 'em':
                run = para.add_run(element.get_text())
                run.italic = True
            
            elif element.name == 'u':
                run = para.add_run(element.get_text())
                run.font.underline = True
            
            elif element.name in ['del', 's', 'strike']:
                run = para.add_run(element.get_text())
                run.font.strike = True
            
            elif element.name == 'sup':
                run = para.add_run(element.get_text())
                run.font.superscript = True
            
            elif element.name == 'sub':
                run = para.add_run(element.get_text())
                run.font.subscript = True
            
            elif element.name == 'code':
                run = para.add_run(element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(199, 37, 78)  # Pink-red
            
            elif element.name == 'span':
                # Check for style attribute (color)
                style = element.get('style', '')
                run = para.add_run(element.get_text())
                
                if 'color:' in style:
                    color_match = re.search(r'color:\s*([^;]+)', style)
                    if color_match:
                        color_value = color_match.group(1).strip()
                        if color_value.startswith('#'):
                            hex_color = color_value.lstrip('#')
                            if len(hex_color) == 6:
                                try:
                                    r = int(hex_color[0:2], 16)
                                    g = int(hex_color[2:4], 16)
                                    b = int(hex_color[4:6], 16)
                                    run.font.color.rgb = RGBColor(r, g, b)
                                except ValueError:
                                    pass
            
            else:
                # Unknown tag - just get text
                para.add_run(element.get_text())

