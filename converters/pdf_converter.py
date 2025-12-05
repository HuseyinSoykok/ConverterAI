"""
PDF converter - handles all PDF conversions
"""
import time
import re
from pathlib import Path
from typing import Optional
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
from PIL import Image
import io

from converters.base import BaseConverter, ConversionResult
from utils.logger import logger


class PDFConverter(BaseConverter):
    """Convert PDF to other formats"""
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        if not text:
            return ''
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))
    
    def _clean_text_for_xml(self, text: str) -> str:
        """Remove XML-incompatible characters (control characters, NULL bytes)"""
        if not text:
            return ''
        import re
        # Remove control characters except newline, tab, carriage return
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        return cleaned
    
    def _fix_word_spacing(self, text: str) -> str:
        """Fix missing spaces between words (common in PDF extraction from presentations)"""
        import re
        if not text:
            return ''
        
        # First pass: Fix obvious concatenations
        # Add space between lowercase and uppercase (camelCase)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Add space after punctuation if followed by letter (not decimal)
        text = re.sub(r'([.])([A-Z])', r'\1 \2', text)
        text = re.sub(r'([,:;!?])([A-Za-z])', r'\1 \2', text)
        
        # Add space after closing parenthesis/bracket if followed by letter
        text = re.sub(r'([\)\]])([A-Za-z])', r'\1 \2', text)
        
        # Add space before opening parenthesis/bracket if preceded by letter
        text = re.sub(r'([a-zA-Z])([\(\[])', r'\1 \2', text)
        
        # Second pass: Fix common concatenated words
        # These patterns are common in PDF extraction errors
        word_boundaries = [
            (r'([a-z])(the)([A-Z])', r'\1 \2 \3'),
            (r'([a-z])(and)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(for)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(with)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(from)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(that)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(this)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(such)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(given)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(find)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(using)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(where)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(which)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(each)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(also)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(into)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(over)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(then)([a-z])', r'\1 \2 \3'),
            (r'([a-z])(when)([a-z])', r'\1 \2 \3'),
            (r'on([A-Z])', r'on \1'),  # "onA" -> "on A"
            (r'in([A-Z])', r'in \1'),  # "inA" -> "in A"
            (r'of([A-Z])', r'of \1'),  # "ofA" -> "of A"
            (r'to([A-Z])', r'to \1'),  # "toA" -> "to A"
            (r'is([A-Z])', r'is \1'),  # "isA" -> "is A"
            (r'by([A-Z])', r'by \1'),  # "byA" -> "by A"
            (r'as([A-Z])', r'as \1'),  # "asA" -> "as A"
        ]
        
        for pattern, replacement in word_boundaries:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # Clean up multiple spaces
        text = re.sub(r' +', ' ', text)
        
        return text
    
    def _fix_concatenated_text(self, text: str) -> str:
        """Advanced fix for heavily concatenated text (pdfplumber extraction issues)"""
        import re
        if not text:
            return ''
        
        # Common English word patterns that get concatenated
        common_words = [
            'the', 'and', 'for', 'with', 'from', 'that', 'this', 'such', 
            'given', 'find', 'using', 'where', 'which', 'each', 'also',
            'into', 'over', 'then', 'when', 'have', 'been', 'will', 'would',
            'could', 'should', 'must', 'function', 'domain', 'source',
            'equation', 'boundary', 'condition', 'method', 'element',
            'finite', 'described', 'strong', 'form', 'unknown', 'potential',
            'temperature', 'heat', 'problem', 'requires', 'twice',
            'differentiate', 'restrictive', 'difficult', 'computationally',
            'notation', 'inner', 'product', 'goal', 'on', 'is', 'by', 'to',
            'an', 'or', 'be', 'as', 'at', 'if', 'of', 'in', 'it'
        ]
        
        for word in common_words:
            # Add space before word if preceded by lowercase letter
            pattern = rf'([a-z])({word})([a-z])'
            text = re.sub(pattern, rf'\1 \2 \3', text, flags=re.IGNORECASE)
        
        # Fix "e.g.," and "i.e.," patterns
        text = re.sub(r'(\w)(e\.g\.,)', r'\1 \2', text)
        text = re.sub(r'(\w)(i\.e\.,)', r'\1 \2', text)
        
        # Fix Greek letters
        text = re.sub(r'([ΩΔ∂∇])([a-zA-Z])', r'\1 \2', text)
        text = re.sub(r'([a-zA-Z])([ΩΔ∂∇])', r'\1 \2', text)
        
        # Clean up multiple spaces
        text = re.sub(r' +', ' ', text)
        
        return text
    
    def convert(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Route to appropriate conversion method
        
        Options:
            use_ocr (bool): Use OCR for text extraction (better for presentations)
            enhance_math (bool): Use AI-enhanced math processing (default: True with OCR)
            use_llm (bool): Use LLM for post-processing (requires API key or Ollama)
            llm_provider (str): LLM provider - 'auto', 'ollama', 'huggingface', 'gemini'
            ocr_lang (str): OCR language code (default: 'eng')
            ocr_dpi (int): OCR resolution DPI multiplier (default: 2)
        """
        output_format = Path(output_file).suffix.lower().lstrip('.')
        use_ocr = options.get('use_ocr', False)
        enhance_math = options.get('enhance_math', True)  # Default to True when OCR is used
        use_llm = options.get('use_llm', False)
        
        start_time = time.time()
        
        # Validate files
        error = self._validate_files(input_file, output_file)
        if error:
            return self._create_error_result(input_file, error, 'pdf', output_format)
        
        try:
            if output_format in ['docx', 'doc']:
                if use_ocr:
                    result = self._pdf_to_docx_ocr(input_file, output_file, **options)
                else:
                    result = self._pdf_to_docx(input_file, output_file, **options)
            elif output_format in ['md', 'markdown']:
                if use_ocr:
                    if use_llm:
                        result = self._pdf_to_markdown_llm(input_file, output_file, **options)
                    elif enhance_math:
                        result = self._pdf_to_markdown_math_ocr(input_file, output_file, **options)
                    else:
                        result = self._pdf_to_markdown_ocr(input_file, output_file, **options)
                else:
                    result = self._pdf_to_markdown(input_file, output_file, **options)
            elif output_format in ['html', 'htm']:
                if use_ocr:
                    result = self._pdf_to_html_ocr(input_file, output_file, **options)
                else:
                    result = self._pdf_to_html(input_file, output_file, **options)
            else:
                return self._create_error_result(
                    input_file, 
                    f"Unsupported output format: {output_format}",
                    'pdf',
                    output_format
                )
            
            processing_time = time.time() - start_time
            result.processing_time = processing_time
            return result
            
        except Exception as e:
            logger.error(f"PDF conversion failed: {e}", exc_info=True)
            return self._create_error_result(input_file, str(e), 'pdf', output_format)
    
    def _pdf_to_docx(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert PDF to DOCX with improved text extraction and formatting"""
        logger.info(f"Converting PDF to DOCX: {input_file} -> {output_file}")
        
        doc = Document()
        warnings = []
        
        try:
            # Open PDF with PyMuPDF for better text extraction
            pdf_document = fitz.open(input_file)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Try table extraction first with pdfplumber
                try:
                    import pdfplumber
                    with pdfplumber.open(input_file) as pdf_plumber:
                        plumber_page = pdf_plumber.pages[page_num]
                        tables = plumber_page.extract_tables()
                        
                        if tables:
                            for table_data in tables:
                                if table_data and len(table_data) > 0:
                                    # Add table to document
                                    num_rows = len(table_data)
                                    num_cols = max(len(row) for row in table_data)
                                    
                                    table = doc.add_table(rows=num_rows, cols=num_cols)
                                    table.style = 'Table Grid'
                                    
                                    for i, row in enumerate(table_data):
                                        for j, cell in enumerate(row):
                                            if j < num_cols and cell:
                                                table.rows[i].cells[j].text = str(cell).strip()
                                    
                                    # Add spacing after table
                                    doc.add_paragraph()
                except Exception as e:
                    warnings.append(f"Table extraction failed: {e}")
                
                # Extract text with formatting
                text_dict = page.get_text("dict")
                blocks = text_dict.get("blocks", [])
                
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        block_text = []
                        max_font_size = 0
                        is_bold = False
                        
                        for line in block.get("lines", []):
                            line_text = ""
                            line_spans = line.get("spans", [])
                            
                            for span_idx, span in enumerate(line_spans):
                                span_text = span.get("text", "")
                                
                                # Smart span joining - preserve leading spaces in span
                                if span_idx > 0:
                                    # If current span starts with space, keep it
                                    # If not, add space between spans (unless previous ended with hyphen)
                                    if not span_text.startswith(' ') and line_text and not line_text.endswith('-'):
                                        # Check if we need a space
                                        if line_text[-1:].isalnum() and span_text[:1].isalnum():
                                            line_text += " "
                                
                                line_text += span_text
                                
                                # Track font properties
                                font_size = span.get("size", 12)
                                if font_size > max_font_size:
                                    max_font_size = font_size
                                
                                # Check if bold
                                font_flags = span.get("flags", 0)
                                if font_flags & 2**4:  # Bold flag
                                    is_bold = True
                            
                            if line_text.strip():
                                block_text.append(line_text.strip())
                        
                        if block_text:
                            full_text = " ".join(block_text)
                            
                            if full_text.strip():
                                # Clean text for XML compatibility and fix spacing
                                full_text = self._clean_text_for_xml(full_text)
                                full_text = self._fix_word_spacing(full_text)
                                
                                # Detect if it's a heading based on font size and properties
                                para = doc.add_paragraph()
                                
                                # Smart heading detection
                                if max_font_size > 18 or (max_font_size > 16 and is_bold):
                                    para.style = 'Heading 1'
                                elif max_font_size > 14 or (max_font_size > 12 and is_bold and len(full_text) < 100):
                                    para.style = 'Heading 2'
                                elif max_font_size > 12 and is_bold and len(full_text) < 80:
                                    para.style = 'Heading 3'
                                else:
                                    para.style = 'Normal'
                                
                                # Add text to paragraph
                                run = para.add_run(full_text)
                                if is_bold and para.style == 'Normal':
                                    run.bold = True
                                
                                # Add spacing after paragraph for better readability
                                para.paragraph_format.space_after = Pt(6)
                    
                    elif block.get("type") == 1:  # Image block
                        try:
                            # Extract and add image
                            img = block.get("image")
                            if img:
                                xref = block.get("xref")
                                if xref:
                                    base_image = pdf_document.extract_image(xref)
                                    image_bytes = base_image["image"]
                                    
                                    # Save temporarily and add to document
                                    image_stream = io.BytesIO(image_bytes)
                                    doc.add_picture(image_stream, width=Inches(5.0))
                        except Exception as e:
                            warnings.append(f"Could not extract image: {e}")
                
                # Page break after each page (except last)
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
            
            num_pages = len(pdf_document)
            pdf_document.close()
            
            # Save document
            doc.save(output_file)
            
            logger.info(f"Successfully converted PDF to DOCX: {output_file}")
            return self._create_success_result(
                input_file, 
                output_file, 
                'pdf', 
                'docx',
                warnings=warnings,
                metadata={'pages': num_pages}
            )
            
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {e}")
            raise
    
    def _pdf_to_markdown(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert PDF to Markdown with enhanced table, list, and formatting support"""
        logger.info(f"Converting PDF to Markdown: {input_file} -> {output_file}")
        
        markdown_content = []
        warnings = []
        
        try:
            # Use both PyMuPDF (for formatting) and pdfplumber (for tables)
            doc = fitz.open(input_file)
            num_pages = len(doc)
            
            for page_num in range(num_pages):
                page = doc[page_num]
                
                # Add page header with proper spacing
                if page_num > 0:
                    markdown_content.append("\n\n---\n\n")
                
                markdown_content.append(f"## Page {page_num + 1}\n\n")
                
                # STEP 1: Extract tables first with pdfplumber
                tables_extracted = []
                table_bboxes = []
                
                try:
                    with pdfplumber.open(input_file) as pdf_plumber:
                        plumber_page = pdf_plumber.pages[page_num]
                        tables = plumber_page.extract_tables()
                        
                        if tables:
                            for table_data in tables:
                                if table_data and len(table_data) > 0:
                                    # Convert table to markdown
                                    md_table = self._table_to_markdown(table_data)
                                    tables_extracted.append(md_table)
                                    
                                    # Try to get table bbox (approximate)
                                    try:
                                        table_obj = plumber_page.find_tables()
                                        if table_obj:
                                            for t in table_obj:
                                                table_bboxes.append(t.bbox)
                                    except:
                                        pass
                except Exception as e:
                    warnings.append(f"Table extraction failed on page {page_num + 1}: {e}")
                
                # STEP 2: Extract text blocks with formatting
                blocks = page.get_text("dict")["blocks"]
                
                # Calculate average font size for the page
                all_font_sizes = []
                for block in blocks:
                    if block.get("type") == 0:
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                all_font_sizes.append(span.get("size", 12))
                
                avg_font_size = sum(all_font_sizes) / len(all_font_sizes) if all_font_sizes else 12
                
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        # Check if block overlaps with a table (skip if yes)
                        block_bbox = block.get("bbox", [0, 0, 0, 0])
                        is_in_table = False
                        
                        for table_bbox in table_bboxes:
                            if self._bbox_overlap(block_bbox, table_bbox):
                                is_in_table = True
                                break
                        
                        if is_in_table:
                            continue
                        
                        # Process lines in block
                        for line in block.get("lines", []):
                            line_text_parts = []
                            line_font_size = 0
                            is_bold = False
                            is_italic = False
                            is_monospace = False
                            
                            for span in line.get("spans", []):
                                span_text = span.get("text", "").strip()
                                if not span_text:
                                    continue
                                
                                # Font properties
                                font_size = span.get("size", 12)
                                font_flags = span.get("flags", 0)
                                font_name = span.get("font", "").lower()
                                
                                if font_size > line_font_size:
                                    line_font_size = font_size
                                
                                # Check bold (bit 4 = 16)
                                span_bold = bool(font_flags & (1 << 4))
                                # Check italic (bit 1 = 2)
                                span_italic = bool(font_flags & (1 << 1))
                                # Check monospace fonts
                                span_monospace = any(mono in font_name for mono in ['courier', 'mono', 'consolas', 'menlo'])
                                
                                # Apply formatting
                                formatted_text = span_text
                                
                                if span_monospace:
                                    formatted_text = f"`{formatted_text}`"
                                    is_monospace = True
                                elif span_bold and span_italic:
                                    formatted_text = f"***{formatted_text}***"
                                elif span_bold:
                                    formatted_text = f"**{formatted_text}**"
                                    is_bold = True
                                elif span_italic:
                                    formatted_text = f"*{formatted_text}*"
                                    is_italic = True
                                
                                line_text_parts.append(formatted_text)
                            
                            if not line_text_parts:
                                continue
                            
                            full_line = " ".join(line_text_parts)
                            
                            # Smart content type detection
                            
                            # 1. Heading detection (based on font size)
                            if line_font_size > avg_font_size * 1.5:
                                heading_level = 1
                            elif line_font_size > avg_font_size * 1.3:
                                heading_level = 2
                            elif line_font_size > avg_font_size * 1.15:
                                heading_level = 3
                            else:
                                heading_level = 0
                            
                            # Additional heading indicators
                            if heading_level == 0:
                                # All caps + short = heading
                                if len(full_line) < 100 and full_line.upper() == full_line and len(full_line) > 3:
                                    heading_level = 3
                                    full_line = full_line.title()
                                # Ends with colon + short = subheading
                                elif len(full_line) < 80 and full_line.endswith(':'):
                                    heading_level = 4
                            
                            if heading_level > 0:
                                # Remove markdown formatting from headings
                                clean_line = full_line.replace('**', '').replace('*', '').replace('`', '')
                                markdown_content.append(f"\n{'#' * heading_level} {clean_line}\n\n")
                                continue
                            
                            # 2. List detection (bullet or numbered)
                            stripped_line = full_line.lstrip()
                            
                            # Bullet list detection
                            if stripped_line and stripped_line[0] in ['•', '·', '◦', '▪', '▫', '-', '–', '—']:
                                list_text = stripped_line[1:].strip()
                                markdown_content.append(f"- {list_text}\n")
                                continue
                            
                            # Numbered list detection (1., a., i., etc.)
                            import re
                            numbered_match = re.match(r'^(\d+|[a-z]|[ivxlcdm]+)[\.\)]\s+(.+)', stripped_line, re.IGNORECASE)
                            if numbered_match:
                                list_text = numbered_match.group(2)
                                markdown_content.append(f"1. {list_text}\n")
                                continue
                            
                            # 3. Code block detection (multiple monospace spans)
                            if is_monospace or full_line.count('`') > 2:
                                # Remove inline code markers for code block
                                code_line = full_line.replace('`', '')
                                # Check if previous line was also code
                                if markdown_content and markdown_content[-1].startswith('    '):
                                    markdown_content.append(f"    {code_line}\n")
                                else:
                                    markdown_content.append(f"\n    {code_line}\n")
                                continue
                            
                            # 4. Regular paragraph
                            # Handle hyphenation
                            if full_line.endswith('-'):
                                markdown_content.append(full_line[:-1])  # Remove hyphen
                            else:
                                markdown_content.append(full_line)
                                
                                # Add proper line break
                                if not full_line.endswith(('.', '!', '?', ':', ';')):
                                    markdown_content.append(" ")
                                else:
                                    markdown_content.append("\n\n")
                
                # STEP 3: Add extracted tables at the end of page content
                if tables_extracted:
                    markdown_content.append("\n\n")
                    for table_md in tables_extracted:
                        markdown_content.append(table_md)
                        markdown_content.append("\n\n")
            
            doc.close()
            
            # Clean up the final content
            final_content = ''.join(markdown_content)
            
            # Fix word spacing issues from PDF extraction (use advanced fix)
            final_content = self._fix_concatenated_text(final_content)
            final_content = self._fix_word_spacing(final_content)
            
            # Remove excessive blank lines (more than 2)
            final_content = re.sub(r'\n{4,}', '\n\n\n', final_content)
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8-sig') as f:
                f.write(final_content)
            
            logger.info(f"Successfully converted PDF to Markdown: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'pdf',
                'markdown',
                warnings=warnings,
                metadata={'pages': num_pages, 'tables': len(tables_extracted)}
            )
            
        except Exception as e:
            logger.error(f"PDF to Markdown conversion failed: {e}")
            raise
    
    def _table_to_markdown(self, table_data):
        """Convert table data to Markdown table format"""
        if not table_data or len(table_data) == 0:
            return ""
        
        md_lines = []
        
        # Get max columns
        max_cols = max(len(row) for row in table_data)
        
        # Process header row
        if len(table_data) > 0:
            header = table_data[0]
            header_cells = []
            for i in range(max_cols):
                cell = header[i] if i < len(header) else ""
                # Clean cell content
                cell_text = str(cell).strip().replace('\n', ' ').replace('|', '\\|')
                header_cells.append(cell_text)
            
            md_lines.append("| " + " | ".join(header_cells) + " |")
            
            # Separator row
            md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
        
        # Process data rows
        for row in table_data[1:]:
            row_cells = []
            for i in range(max_cols):
                cell = row[i] if i < len(row) else ""
                # Clean cell content
                cell_text = str(cell).strip().replace('\n', ' ').replace('|', '\\|')
                row_cells.append(cell_text)
            
            md_lines.append("| " + " | ".join(row_cells) + " |")
        
        return "\n".join(md_lines)
    
    def _bbox_overlap(self, bbox1, bbox2, threshold=0.5):
        """Check if two bounding boxes overlap significantly"""
        if not bbox1 or not bbox2:
            return False
        
        # bbox format: (x0, y0, x1, y1)
        x_overlap = max(0, min(bbox1[2], bbox2[2]) - max(bbox1[0], bbox2[0]))
        y_overlap = max(0, min(bbox1[3], bbox2[3]) - max(bbox1[1], bbox2[1]))
        
        overlap_area = x_overlap * y_overlap
        bbox1_area = (bbox1[2] - bbox1[0]) * (bbox1[3] - bbox1[1])
        
        if bbox1_area == 0:
            return False
        
        overlap_ratio = overlap_area / bbox1_area
        return overlap_ratio > threshold
    
    def _pdf_to_html(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert PDF to HTML with enhanced styling and readability"""
        logger.info(f"Converting PDF to HTML: {input_file} -> {output_file}")
        
        html_parts = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            '<meta charset="UTF-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
            '<title>Converted PDF</title>',
            '<style>',
            '* { box-sizing: border-box; }',
            'body {',
            '    font-family: "Segoe UI", Arial, sans-serif;',
            '    margin: 0;',
            '    padding: 40px;',
            '    line-height: 1.8;',
            '    color: #333;',
            '    background-color: #f5f5f5;',
            '}',
            '.container {',
            '    max-width: 900px;',
            '    margin: 0 auto;',
            '    background: white;',
            '    padding: 40px;',
            '    box-shadow: 0 2px 8px rgba(0,0,0,0.1);',
            '}',
            '.page {',
            '    margin-bottom: 60px;',
            '    padding-bottom: 40px;',
            '    border-bottom: 2px solid #e0e0e0;',
            '}',
            '.page:last-child { border-bottom: none; }',
            'h1, h2, h3, h4, h5, h6 {',
            '    margin-top: 24px;',
            '    margin-bottom: 16px;',
            '    font-weight: 600;',
            '    line-height: 1.25;',
            '}',
            'h1 {',
            '    font-size: 2em;',
            '    color: #1a1a1a;',
            '    border-bottom: 3px solid #0066cc;',
            '    padding-bottom: 12px;',
            '}',
            'h2 {',
            '    font-size: 1.5em;',
            '    color: #2a2a2a;',
            '    border-bottom: 2px solid #e0e0e0;',
            '    padding-bottom: 8px;',
            '}',
            'h3 { font-size: 1.25em; color: #333; }',
            'h4 { font-size: 1.1em; color: #444; }',
            'p {',
            '    margin: 16px 0;',
            '    text-align: justify;',
            '}',
            'ul, ol {',
            '    margin: 16px 0;',
            '    padding-left: 32px;',
            '}',
            'li {',
            '    margin: 8px 0;',
            '    line-height: 1.6;',
            '}',
            'ul li {',
            '    list-style-type: disc;',
            '}',
            'table {',
            '    border-collapse: collapse;',
            '    width: 100%;',
            '    margin: 24px 0;',
            '    background-color: white;',
            '    box-shadow: 0 1px 3px rgba(0,0,0,0.1);',
            '}',
            'table, th, td {',
            '    border: 1px solid #ddd;',
            '}',
            'th, td {',
            '    padding: 14px 16px;',
            '    text-align: left;',
            '}',
            'th {',
            '    background-color: #0066cc;',
            '    color: white;',
            '    font-weight: 600;',
            '    text-transform: uppercase;',
            '    font-size: 0.9em;',
            '    letter-spacing: 0.5px;',
            '}',
            'tr:nth-child(even) { background-color: #f9f9f9; }',
            'tr:hover { background-color: #f0f0f0; }',
            'img {',
            '    max-width: 100%;',
            '    height: auto;',
            '    margin: 20px 0;',
            '    border-radius: 4px;',
            '    box-shadow: 0 2px 4px rgba(0,0,0,0.1);',
            '}',
            'code {',
            '    background-color: #f4f4f4;',
            '    padding: 2px 6px;',
            '    border-radius: 3px;',
            '    font-family: "Courier New", monospace;',
            '    font-size: 0.9em;',
            '}',
            'pre {',
            '    background-color: #f4f4f4;',
            '    padding: 16px;',
            '    border-radius: 4px;',
            '    overflow-x: auto;',
            '    margin: 20px 0;',
            '}',
            '.page-number {',
            '    color: #666;',
            '    font-size: 0.9em;',
            '    margin-bottom: 20px;',
            '}',
            '@media print {',
            '    body { background: white; }',
            '    .container { box-shadow: none; }',
            '    .page { page-break-after: always; }',
            '}',
            '</style>',
            '</head>',
            '<body>',
            '<div class="container">',
        ]
        
        warnings = []
        
        try:
            # Use pdfplumber for better structured extraction
            with pdfplumber.open(input_file) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    html_parts.append(f'<div class="page">')
                    html_parts.append(f'<div class="page-number">Page {page_num} of {len(pdf.pages)}</div>')
                    
                    # Extract text with layout preservation
                    text = page.extract_text()
                    if text:
                        # Fix word spacing issues (use advanced fix for concatenated text)
                        text = self._fix_concatenated_text(text)
                        text = self._fix_word_spacing(text)
                        
                        # Split into lines for better processing
                        lines = text.split('\n')
                        current_paragraph = []
                        in_list = False
                        
                        for line in lines:
                            line = line.strip()
                            if not line:
                                # Empty line - end current paragraph
                                if current_paragraph:
                                    para_text = ' '.join(current_paragraph)
                                    html_parts.append(f'<p>{self._escape_html(para_text)}</p>')
                                    current_paragraph = []
                                if in_list:
                                    html_parts.append('</ul>')
                                    in_list = False
                                continue
                            
                            # Detect headings by font size and formatting
                            # Check if line is all uppercase and short (likely heading)
                            if len(line) < 100 and line.isupper() and not any(char.isdigit() for char in line[:3]):
                                if current_paragraph:
                                    para_text = ' '.join(current_paragraph)
                                    html_parts.append(f'<p>{self._escape_html(para_text)}</p>')
                                    current_paragraph = []
                                html_parts.append(f'<h2>{self._escape_html(line.title())}</h2>')
                            
                            # Detect subheadings (lines ending with colon)
                            elif len(line) < 80 and line.endswith(':') and not line.startswith(' '):
                                if current_paragraph:
                                    para_text = ' '.join(current_paragraph)
                                    html_parts.append(f'<p>{self._escape_html(para_text)}</p>')
                                    current_paragraph = []
                                html_parts.append(f'<h3>{self._escape_html(line)}</h3>')
                            
                            # Detect bullet points or numbered lists
                            elif line.startswith(('•', '-', '*', '▪', '◦', '▫', '■', '□')) or \
                                 (len(line) > 2 and line[0].isdigit() and line[1] in '.):'):
                                if current_paragraph:
                                    para_text = ' '.join(current_paragraph)
                                    html_parts.append(f'<p>{self._escape_html(para_text)}</p>')
                                    current_paragraph = []
                                if not in_list:
                                    html_parts.append('<ul>')
                                    in_list = True
                                # Remove bullet/number prefix
                                list_text = line.lstrip('•-*▪◦▫■□ ')
                                if line[0].isdigit():
                                    list_text = line.split('.', 1)[1].strip() if '.' in line else line
                                html_parts.append(f'<li>{self._escape_html(list_text)}</li>')
                            
                            # Regular text line
                            else:
                                if in_list:
                                    html_parts.append('</ul>')
                                    in_list = False
                                current_paragraph.append(line)
                        
                        # Close any remaining paragraph or list
                        if current_paragraph:
                            para_text = ' '.join(current_paragraph)
                            html_parts.append(f'<p>{self._escape_html(para_text)}</p>')
                        if in_list:
                            html_parts.append('</ul>')
                    
                    # Extract tables with better formatting
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table and len(table) > 0:
                                html_parts.append('<table>')
                                
                                for row_idx, row in enumerate(table):
                                    html_parts.append('<tr>')
                                    tag = 'th' if row_idx == 0 else 'td'
                                    
                                    for cell in row:
                                        if cell:
                                            # Clean and escape cell content
                                            cell_content = ' '.join(str(cell).split())
                                            cell_content = self._escape_html(cell_content)
                                        else:
                                            cell_content = ''
                                        html_parts.append(f'<{tag}>{cell_content}</{tag}>')
                                    
                                    html_parts.append('</tr>')
                                
                                html_parts.append('</table>')
                    
                    html_parts.append('</div>')
            
            html_parts.extend(['</div>', '</body>', '</html>'])
            
            # Write to file with UTF-8 encoding
            final_html = '\n'.join(html_parts)
            with open(output_file, 'w', encoding='utf-8-sig') as f:
                f.write(final_html)
            
            logger.info(f"Successfully converted PDF to HTML: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'pdf',
                'html',
                warnings=warnings,
                metadata={'pages': len(pdf.pages)}
            )
            
        except Exception as e:
            logger.error(f"PDF to HTML conversion failed: {e}")
            raise

    # ==================== OCR-BASED CONVERSION METHODS ====================
    
    def _setup_tesseract(self):
        """Setup Tesseract OCR path"""
        import pytesseract
        import os
        
        # Common Tesseract installation paths
        tesseract_paths = [
            r'D:\APPS\OCR Tesseract\tesseract.exe',
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            r'/usr/bin/tesseract',
            r'/usr/local/bin/tesseract',
        ]
        
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return True
        
        # Try system PATH
        try:
            import subprocess
            subprocess.run(['tesseract', '--version'], capture_output=True, check=True)
            return True
        except:
            pass
        
        return False
    
    def _pdf_page_to_image(self, page, dpi_multiplier: int = 2) -> Image.Image:
        """Convert a PyMuPDF page to PIL Image"""
        # Create high-resolution pixmap
        matrix = fitz.Matrix(dpi_multiplier, dpi_multiplier)
        pix = page.get_pixmap(matrix=matrix)
        
        # Convert to PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        return img
    
    def _ocr_image(self, img: Image.Image, lang: str = 'eng') -> str:
        """Perform OCR on an image"""
        import pytesseract
        
        # Configure Tesseract for better results
        custom_config = r'--oem 3 --psm 6'
        
        text = pytesseract.image_to_string(img, lang=lang, config=custom_config)
        return text
    
    def _pdf_to_markdown_ocr(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert PDF to Markdown using OCR (better for presentations and scanned documents)"""
        logger.info(f"Converting PDF to Markdown (OCR mode): {input_file} -> {output_file}")
        
        warnings = []
        ocr_lang = options.get('ocr_lang', 'eng')
        dpi_multiplier = options.get('ocr_dpi', 2)
        
        # Setup Tesseract
        if not self._setup_tesseract():
            warnings.append("Tesseract not found, OCR may not work properly")
        
        try:
            import pytesseract
            doc = fitz.open(input_file)
            markdown_content = []
            
            # Get document title from metadata or filename
            metadata = doc.metadata
            title = metadata.get('title', '') if metadata else ''
            if not title:
                title = Path(input_file).stem.replace('_', ' ')
            
            markdown_content.append(f"# {title}\n\n")
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to image
                img = self._pdf_page_to_image(page, dpi_multiplier)
                
                # OCR the image
                try:
                    text = self._ocr_image(img, ocr_lang)
                    
                    if text.strip():
                        # Add page header for multi-page documents
                        if len(doc) > 1:
                            markdown_content.append(f"---\n\n## Page {page_num + 1}\n\n")
                        
                        # Clean up OCR text
                        cleaned_text = self._clean_ocr_text(text)
                        markdown_content.append(cleaned_text)
                        markdown_content.append("\n\n")
                    else:
                        warnings.append(f"Page {page_num + 1}: No text detected via OCR")
                        
                except Exception as e:
                    warnings.append(f"Page {page_num + 1}: OCR failed - {e}")
            
            num_pages = len(doc)
            doc.close()
            
            # Write to file
            final_content = ''.join(markdown_content)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(final_content)
            
            logger.info(f"Successfully converted PDF to Markdown (OCR): {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'pdf',
                'markdown',
                warnings=warnings,
                metadata={'pages': num_pages, 'method': 'ocr'}
            )
            
        except ImportError:
            raise Exception("pytesseract not installed. Install with: pip install pytesseract")
        except Exception as e:
            logger.error(f"PDF to Markdown (OCR) conversion failed: {e}")
            raise
    
    def _clean_ocr_text(self, text: str) -> str:
        """Clean up OCR output text"""
        if not text:
            return ''
        
        # Remove multiple blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Fix common OCR errors
        text = text.replace('|', 'I')  # Common OCR confusion
        
        # Clean up line breaks within paragraphs (keep paragraph breaks)
        lines = text.split('\n')
        cleaned_lines = []
        current_paragraph = []
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                # Empty line indicates paragraph break
                if current_paragraph:
                    cleaned_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                cleaned_lines.append('')
            else:
                current_paragraph.append(stripped)
        
        if current_paragraph:
            cleaned_lines.append(' '.join(current_paragraph))
        
        return '\n\n'.join([l for l in cleaned_lines if l or (cleaned_lines.index(l) > 0 and cleaned_lines[cleaned_lines.index(l)-1])])
    
    def _pdf_to_html_ocr(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert PDF to HTML using OCR"""
        logger.info(f"Converting PDF to HTML (OCR mode): {input_file} -> {output_file}")
        
        warnings = []
        ocr_lang = options.get('ocr_lang', 'eng')
        dpi_multiplier = options.get('ocr_dpi', 2)
        
        # Setup Tesseract
        if not self._setup_tesseract():
            warnings.append("Tesseract not found, OCR may not work properly")
        
        try:
            import pytesseract
            doc = fitz.open(input_file)
            
            # Get document title
            metadata = doc.metadata
            title = metadata.get('title', '') if metadata else ''
            if not title:
                title = Path(input_file).stem.replace('_', ' ')
            
            html_parts = [
                '<!DOCTYPE html>',
                '<html lang="en">',
                '<head>',
                '<meta charset="UTF-8">',
                '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
                f'<title>{self._escape_html(title)}</title>',
                '<style>',
                'body { font-family: Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; }',
                '.page { margin-bottom: 40px; padding-bottom: 20px; border-bottom: 1px solid #ddd; }',
                '.page-header { color: #666; font-size: 14px; margin-bottom: 20px; }',
                'p { line-height: 1.6; margin-bottom: 1em; }',
                '</style>',
                '</head>',
                '<body>',
                f'<h1>{self._escape_html(title)}</h1>',
            ]
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to image
                img = self._pdf_page_to_image(page, dpi_multiplier)
                
                # OCR the image
                try:
                    text = self._ocr_image(img, ocr_lang)
                    
                    html_parts.append(f'<div class="page">')
                    if len(doc) > 1:
                        html_parts.append(f'<div class="page-header">Page {page_num + 1}</div>')
                    
                    if text.strip():
                        # Convert text to HTML paragraphs
                        paragraphs = text.strip().split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                clean_para = ' '.join(para.split())
                                html_parts.append(f'<p>{self._escape_html(clean_para)}</p>')
                    else:
                        warnings.append(f"Page {page_num + 1}: No text detected via OCR")
                    
                    html_parts.append('</div>')
                    
                except Exception as e:
                    warnings.append(f"Page {page_num + 1}: OCR failed - {e}")
            
            num_pages = len(doc)
            doc.close()
            
            html_parts.extend(['</body>', '</html>'])
            
            # Write to file
            final_html = '\n'.join(html_parts)
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(final_html)
            
            logger.info(f"Successfully converted PDF to HTML (OCR): {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'pdf',
                'html',
                warnings=warnings,
                metadata={'pages': num_pages, 'method': 'ocr'}
            )
            
        except ImportError:
            raise Exception("pytesseract not installed. Install with: pip install pytesseract")
        except Exception as e:
            logger.error(f"PDF to HTML (OCR) conversion failed: {e}")
            raise
    
    def _pdf_to_docx_ocr(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert PDF to DOCX using OCR"""
        logger.info(f"Converting PDF to DOCX (OCR mode): {input_file} -> {output_file}")
        
        warnings = []
        ocr_lang = options.get('ocr_lang', 'eng')
        dpi_multiplier = options.get('ocr_dpi', 2)
        
        # Setup Tesseract
        if not self._setup_tesseract():
            warnings.append("Tesseract not found, OCR may not work properly")
        
        try:
            import pytesseract
            doc = Document()
            pdf_doc = fitz.open(input_file)
            
            # Get document title
            metadata = pdf_doc.metadata
            title = metadata.get('title', '') if metadata else ''
            if not title:
                title = Path(input_file).stem.replace('_', ' ')
            
            # Add title
            title_para = doc.add_heading(title, level=0)
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                
                # Convert page to image
                img = self._pdf_page_to_image(page, dpi_multiplier)
                
                # OCR the image
                try:
                    text = self._ocr_image(img, ocr_lang)
                    
                    if len(pdf_doc) > 1:
                        doc.add_heading(f'Page {page_num + 1}', level=1)
                    
                    if text.strip():
                        # Add text paragraphs
                        paragraphs = text.strip().split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                clean_para = ' '.join(para.split())
                                clean_para = self._clean_text_for_xml(clean_para)
                                doc.add_paragraph(clean_para)
                    else:
                        warnings.append(f"Page {page_num + 1}: No text detected via OCR")
                    
                    # Add page break except for last page
                    if page_num < len(pdf_doc) - 1:
                        doc.add_page_break()
                        
                except Exception as e:
                    warnings.append(f"Page {page_num + 1}: OCR failed - {e}")
            
            num_pages = len(pdf_doc)
            pdf_doc.close()
            
            # Save document
            doc.save(output_file)
            
            logger.info(f"Successfully converted PDF to DOCX (OCR): {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'pdf',
                'docx',
                warnings=warnings,
                metadata={'pages': num_pages, 'method': 'ocr'}
            )
            
        except ImportError:
            raise Exception("pytesseract not installed. Install with: pip install pytesseract")
        except Exception as e:
            logger.error(f"PDF to DOCX (OCR) conversion failed: {e}")
            raise
    
    def detect_if_ocr_needed(self, input_file: str) -> dict:
        """Analyze PDF and suggest whether OCR should be used
        
        Returns:
            dict with 'recommended': bool, 'reason': str, 'confidence': float
        """
        try:
            doc = fitz.open(input_file)
            
            # Check 1: Page aspect ratio (presentation = 16:9 or 4:3)
            if len(doc) > 0:
                page = doc[0]
                rect = page.rect
                aspect_ratio = rect.width / rect.height if rect.height > 0 else 1
                
                is_presentation = 1.3 < aspect_ratio < 2.0  # 4:3 = 1.33, 16:9 = 1.78
                
                # Check 2: Sample text extraction quality
                sample_text = ""
                for i in range(min(3, len(doc))):
                    sample_text += doc[i].get_text()
                
                # Check for concatenated words (no spaces between capitals)
                concatenation_pattern = re.compile(r'[a-z][A-Z]')
                concatenation_count = len(concatenation_pattern.findall(sample_text))
                
                # Check for very long words (likely concatenated)
                words = sample_text.split()
                long_words = [w for w in words if len(w) > 25]
                
                doc.close()
                
                # Decision logic
                if is_presentation and (concatenation_count > 10 or len(long_words) > 5):
                    return {
                        'recommended': True,
                        'reason': 'Presentation format with text extraction issues detected',
                        'confidence': 0.9
                    }
                elif concatenation_count > 20 or len(long_words) > 10:
                    return {
                        'recommended': True,
                        'reason': 'Significant text extraction issues detected',
                        'confidence': 0.8
                    }
                elif is_presentation:
                    return {
                        'recommended': True,
                        'reason': 'Presentation format detected (OCR often produces better results)',
                        'confidence': 0.6
                    }
                else:
                    return {
                        'recommended': False,
                        'reason': 'Standard document, text extraction should work well',
                        'confidence': 0.7
                    }
                    
        except Exception as e:
            return {
                'recommended': False,
                'reason': f'Could not analyze PDF: {e}',
                'confidence': 0.0
            }
    
    def _pdf_to_markdown_math_ocr(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """
        Convert PDF to Markdown using OCR with enhanced math processing.
        This method produces high-quality output with proper LaTeX math formatting.
        Best for: academic papers, presentations with equations, technical documents.
        """
        logger.info(f"Converting PDF to Markdown (Math-Enhanced OCR): {input_file} -> {output_file}")
        
        warnings = []
        ocr_lang = options.get('ocr_lang', 'eng')
        dpi_multiplier = options.get('ocr_dpi', 3)  # Higher DPI for better math recognition
        
        # Setup Tesseract
        if not self._setup_tesseract():
            warnings.append("Tesseract not found, OCR may not work properly")
        
        try:
            import pytesseract
            from ai.math_ocr_processor import MathOCRProcessor
            
            doc = fitz.open(input_file)
            
            # Get document title from metadata or filename
            metadata = doc.metadata
            title = metadata.get('title', '') if metadata else ''
            if not title:
                title = Path(input_file).stem.replace('_', ' ')
            
            # Process all pages with high-resolution OCR
            all_page_texts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # High-resolution rendering for better OCR
                matrix = fitz.Matrix(dpi_multiplier, dpi_multiplier)
                pix = page.get_pixmap(matrix=matrix)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # OCR with settings optimized for academic content
                custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
                
                try:
                    text = pytesseract.image_to_string(img, lang=ocr_lang, config=custom_config)
                    
                    if text.strip():
                        all_page_texts.append(f"## Page {page_num + 1}\n\n{text}")
                    else:
                        warnings.append(f"Page {page_num + 1}: No text detected via OCR")
                        
                except Exception as e:
                    warnings.append(f"Page {page_num + 1}: OCR failed - {e}")
            
            num_pages = len(doc)
            doc.close()
            
            # Combine all pages
            raw_ocr_text = '\n\n---\n\n'.join(all_page_texts)
            
            # Apply math-enhanced processing
            processor = MathOCRProcessor()
            processed_content = processor.process(raw_ocr_text, title=title)
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(processed_content)
            
            logger.info(f"Successfully converted PDF to Markdown (Math-Enhanced OCR): {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'pdf',
                'markdown',
                warnings=warnings,
                metadata={
                    'pages': num_pages, 
                    'method': 'math_enhanced_ocr',
                    'dpi': dpi_multiplier * 72,
                    'lang': ocr_lang
                }
            )
            
        except ImportError as e:
            logger.warning(f"Math processor not available: {e}")
            # Fallback to regular OCR
            return self._pdf_to_markdown_ocr(input_file, output_file, **options)
        except Exception as e:
            logger.error(f"PDF to Markdown (Math-Enhanced OCR) conversion failed: {e}")
            raise

    def _pdf_to_markdown_llm(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """
        Convert PDF to Markdown using OCR + LLM post-processing.
        This produces the highest quality output with proper LaTeX math formatting.
        
        Requires one of:
        - Ollama running locally (free, recommended)
        - HuggingFace API key (free tier)
        - Google Gemini API key (free tier)
        """
        logger.info(f"Converting PDF to Markdown (LLM-Enhanced): {input_file} -> {output_file}")
        
        warnings = []
        ocr_lang = options.get('ocr_lang', 'eng')
        dpi_multiplier = options.get('ocr_dpi', 3)  # Higher DPI for better OCR
        llm_provider = options.get('llm_provider', 'auto')
        
        # Setup Tesseract
        if not self._setup_tesseract():
            warnings.append("Tesseract not found, OCR may not work properly")
        
        try:
            import pytesseract
            from ai.llm_post_processor import LLMPostProcessor
            
            # Initialize LLM processor
            llm_config = {
                'huggingface_api_key': options.get('huggingface_api_key'),
                'google_api_key': options.get('google_api_key'),
                'ollama_model': options.get('ollama_model'),
                'ollama_host': options.get('ollama_host'),
            }
            llm_processor = LLMPostProcessor(provider=llm_provider, **llm_config)
            
            if not llm_processor.is_available():
                warnings.append("No LLM provider available, falling back to rule-based processing")
                return self._pdf_to_markdown_math_ocr(input_file, output_file, **options)
            
            logger.info(f"Using LLM provider: {llm_processor.provider_name}")
            
            doc = fitz.open(input_file)
            
            # Get document title
            metadata = doc.metadata
            title = metadata.get('title', '') if metadata else ''
            if not title:
                title = Path(input_file).stem.replace('_', ' ')
            
            # Process all pages with high-resolution OCR
            all_page_texts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # High-resolution rendering
                matrix = fitz.Matrix(dpi_multiplier, dpi_multiplier)
                pix = page.get_pixmap(matrix=matrix)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # OCR with optimized settings
                custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
                
                try:
                    text = pytesseract.image_to_string(img, lang=ocr_lang, config=custom_config)
                    
                    if text.strip():
                        all_page_texts.append(f"=== Page {page_num + 1} ===\n\n{text}")
                    else:
                        warnings.append(f"Page {page_num + 1}: No text detected via OCR")
                        
                except Exception as e:
                    warnings.append(f"Page {page_num + 1}: OCR failed - {e}")
            
            num_pages = len(doc)
            doc.close()
            
            # Combine all pages
            raw_ocr_text = '\n\n'.join(all_page_texts)
            
            # Process with LLM
            logger.info("Processing with LLM (this may take a moment)...")
            processed_content, llm_metadata = llm_processor.process_math_document(raw_ocr_text)
            
            # Add title if not present
            if not processed_content.startswith('#'):
                processed_content = f"# {title}\n\n{processed_content}"
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(processed_content)
            
            logger.info(f"Successfully converted PDF to Markdown (LLM-Enhanced): {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'pdf',
                'markdown',
                warnings=warnings,
                metadata={
                    'pages': num_pages,
                    'method': 'llm_enhanced_ocr',
                    'llm_provider': llm_metadata.get('provider'),
                    'llm_chunks': llm_metadata.get('chunks', 1),
                    'dpi': dpi_multiplier * 72,
                    'lang': ocr_lang
                }
            )
            
        except ImportError as e:
            logger.warning(f"LLM processor not available: {e}")
            return self._pdf_to_markdown_math_ocr(input_file, output_file, **options)
        except Exception as e:
            logger.error(f"PDF to Markdown (LLM-Enhanced) conversion failed: {e}")
            raise
