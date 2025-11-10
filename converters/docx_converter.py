"""
DOCX converter - handles all DOCX conversions
"""
import time
from pathlib import Path
from typing import Optional
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
import markdown
from markdown.extensions.tables import TableExtension
from markdown.extensions.fenced_code import FencedCodeExtension
from bs4 import BeautifulSoup

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, Table as RLTable, TableStyle, PageBreak
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

from converters.base import BaseConverter, ConversionResult
from utils.logger import logger


class DOCXConverter(BaseConverter):
    """Convert DOCX to other formats"""
    
    def convert(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Route to appropriate conversion method"""
        output_format = Path(output_file).suffix.lower().lstrip('.')
        
        start_time = time.time()
        
        # Validate files
        error = self._validate_files(input_file, output_file)
        if error:
            return self._create_error_result(input_file, error, 'docx', output_format)
        
        try:
            if output_format == 'pdf':
                result = self._docx_to_pdf(input_file, output_file, **options)
            elif output_format in ['md', 'markdown']:
                result = self._docx_to_markdown(input_file, output_file, **options)
            elif output_format in ['html', 'htm']:
                result = self._docx_to_html(input_file, output_file, **options)
            else:
                return self._create_error_result(
                    input_file,
                    f"Unsupported output format: {output_format}",
                    'docx',
                    output_format
                )
            
            processing_time = time.time() - start_time
            result.processing_time = processing_time
            return result
            
        except Exception as e:
            logger.error(f"DOCX conversion failed: {e}", exc_info=True)
            return self._create_error_result(input_file, str(e), 'docx', output_format)
    
    def _docx_to_pdf(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert DOCX to PDF"""
        logger.info(f"Converting DOCX to PDF: {input_file} -> {output_file}")
        
        try:
            # First convert to HTML, then to PDF for better cross-platform support
            html_content = self._extract_docx_as_html(input_file)
            
            # Try WeasyPrint first, fallback to reportlab
            try:
                from weasyprint import HTML
                HTML(string=html_content).write_pdf(output_file)
                logger.info("Used WeasyPrint for DOCX to PDF conversion")
            except (ImportError, OSError) as e:
                # Fallback to reportlab with enhanced formatting
                logger.warning(f"WeasyPrint not available ({e}), using reportlab fallback")
                
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Create PDF with enhanced formatting
                doc_pdf = SimpleDocTemplate(
                    output_file, 
                    pagesize=letter,
                    topMargin=0.75*inch,
                    bottomMargin=0.75*inch,
                    leftMargin=0.75*inch,
                    rightMargin=0.75*inch
                )
                
                # Get and customize styles
                styles = getSampleStyleSheet()
                
                # Enhanced Heading styles
                styles.add(ParagraphStyle(
                    name='CustomH1',
                    parent=styles['Heading1'],
                    fontSize=24,
                    textColor=colors.HexColor('#1a1a1a'),
                    spaceAfter=16,
                    spaceBefore=12,
                    fontName='Helvetica-Bold',
                    leading=28
                ))
                
                styles.add(ParagraphStyle(
                    name='CustomH2',
                    parent=styles['Heading2'],
                    fontSize=20,
                    textColor=colors.HexColor('#2d2d2d'),
                    spaceAfter=14,
                    spaceBefore=10,
                    fontName='Helvetica-Bold',
                    leading=24
                ))
                
                styles.add(ParagraphStyle(
                    name='CustomH3',
                    parent=styles['Heading3'],
                    fontSize=16,
                    textColor=colors.HexColor('#404040'),
                    spaceAfter=12,
                    spaceBefore=8,
                    fontName='Helvetica-Bold',
                    leading=20
                ))
                
                styles.add(ParagraphStyle(
                    name='CustomH4',
                    parent=styles['Heading4'],
                    fontSize=14,
                    textColor=colors.HexColor('#555555'),
                    spaceAfter=10,
                    spaceBefore=6,
                    fontName='Helvetica-Bold',
                    leading=18
                ))
                
                # Code block style
                styles.add(ParagraphStyle(
                    name='CodeBlock',
                    parent=styles['Code'],
                    fontName='Courier',
                    fontSize=9,
                    textColor=colors.HexColor('#2d2d2d'),
                    backColor=colors.HexColor('#f5f5f5'),
                    borderColor=colors.HexColor('#dddddd'),
                    borderWidth=1,
                    borderPadding=8,
                    leftIndent=20,
                    rightIndent=20,
                    spaceAfter=12,
                    spaceBefore=12,
                    leading=11
                ))
                
                # List item style
                styles.add(ParagraphStyle(
                    name='ListItem',
                    parent=styles['Normal'],
                    fontSize=11,
                    leftIndent=25,
                    spaceAfter=6,
                    bulletIndent=10
                ))
                
                # Enhanced body text
                styles.add(ParagraphStyle(
                    name='EnhancedBody',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.HexColor('#333333'),
                    spaceAfter=8,
                    leading=15,
                    alignment=0
                ))
                
                story = []
                
                # Process HTML elements with proper formatting
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'code', 'ul', 'ol', 'table']):
                    try:
                        # Headings
                        if element.name == 'h1':
                            text = element.get_text().strip()
                            if text:
                                story.append(RLParagraph(text, styles['CustomH1']))
                                story.append(Spacer(1, 0.2 * inch))
                        
                        elif element.name == 'h2':
                            text = element.get_text().strip()
                            if text:
                                story.append(RLParagraph(text, styles['CustomH2']))
                                story.append(Spacer(1, 0.15 * inch))
                        
                        elif element.name == 'h3':
                            text = element.get_text().strip()
                            if text:
                                story.append(RLParagraph(text, styles['CustomH3']))
                                story.append(Spacer(1, 0.12 * inch))
                        
                        elif element.name in ['h4', 'h5', 'h6']:
                            text = element.get_text().strip()
                            if text:
                                story.append(RLParagraph(text, styles['CustomH4']))
                                story.append(Spacer(1, 0.1 * inch))
                        
                        # Code blocks
                        elif element.name == 'pre':
                            code_text = element.get_text().strip()
                            if code_text:
                                code_text = code_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(RLParagraph(code_text, styles['CodeBlock']))
                        
                        # Lists
                        elif element.name in ['ul', 'ol']:
                            for li in element.find_all('li', recursive=False):
                                li_text = li.get_text().strip()
                                if li_text:
                                    if element.name == 'ul':
                                        bullet = '•'
                                    else:
                                        bullet = f"{element.find_all('li', recursive=False).index(li) + 1}."
                                    story.append(RLParagraph(f"{bullet} {li_text}", styles['ListItem']))
                        
                        # Tables
                        elif element.name == 'table':
                            table_data = []
                            headers = element.find_all('th')
                            if headers:
                                table_data.append([th.get_text().strip() for th in headers])
                            
                            for tr in element.find_all('tr'):
                                tds = tr.find_all('td')
                                if tds:
                                    table_data.append([td.get_text().strip() for td in tds])
                            
                            if table_data:
                                pdf_table = RLTable(table_data)
                                pdf_table.setStyle(TableStyle([
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4a90e2')),
                                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
                                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                                ]))
                                story.append(pdf_table)
                                story.append(Spacer(1, 0.2 * inch))
                        
                        # Regular paragraphs with inline formatting
                        elif element.name == 'p':
                            para_html = str(element)
                            para_html = para_html.replace('<p>', '').replace('</p>', '')
                            para_html = para_html.replace('<strong>', '<b>').replace('</strong>', '</b>')
                            para_html = para_html.replace('<em>', '<i>').replace('</em>', '</i>')
                            para_html = para_html.replace('<code>', '<font name="Courier" color="#c7254e">').replace('</code>', '</font>')
                            
                            text = para_html.strip()
                            if text and text not in ['', ' ']:
                                story.append(RLParagraph(text, styles['EnhancedBody']))
                                story.append(Spacer(1, 0.05 * inch))
                    
                    except Exception as elem_e:
                        logger.warning(f"Error processing DOCX→PDF element {element.name}: {elem_e}")
                        continue
                
                doc_pdf.build(story)
                logger.info("Used ReportLab with enhanced formatting for DOCX to PDF conversion")
            
            logger.info(f"Successfully converted DOCX to PDF: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'docx',
                'pdf'
            )
            
        except Exception as e:
            logger.error(f"DOCX to PDF conversion failed: {e}")
            raise
    
    def _docx_to_markdown(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert DOCX to Markdown with improved formatting"""
        logger.info(f"Converting DOCX to Markdown: {input_file} -> {output_file}")
        
        doc = Document(input_file)
        markdown_lines = []
        warnings = []
        in_list = False
        
        try:
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    para = Paragraph(element, doc)
                    text = para.text.strip()
                    
                    if not text:
                        # Empty paragraph - add spacing
                        if markdown_lines and markdown_lines[-1] != "":
                            markdown_lines.append("")
                        continue
                    
                    # Detect heading style
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            level = int(level)
                            # Add spacing before heading (except at start)
                            if markdown_lines and markdown_lines[-1] != "":
                                markdown_lines.append("")
                            markdown_lines.append(f"{'#' * level} {text}")
                            markdown_lines.append("")  # Space after heading
                            in_list = False
                        except:
                            markdown_lines.append(text)
                            markdown_lines.append("")
                    
                    # Detect list items
                    elif para.style.name.startswith('List'):
                        if 'Bullet' in para.style.name:
                            markdown_lines.append(f"- {text}")
                            in_list = True
                        elif 'Number' in para.style.name:
                            # Try to extract number or use default
                            markdown_lines.append(f"1. {text}")
                            in_list = True
                        else:
                            markdown_lines.append(text)
                            markdown_lines.append("")
                            in_list = False
                    
                    else:
                        # Regular paragraph
                        if in_list:
                            markdown_lines.append("")  # End list
                            in_list = False
                        
                        # IMPROVED: Check for code blocks (monospace fonts)
                        is_code_block = False
                        if para.runs:
                            # Detect code by font name
                            code_fonts = ['Courier New', 'Consolas', 'Monaco', 'Menlo', 'Courier', 
                                         'Lucida Console', 'Source Code Pro', 'Fira Code']
                            
                            # Check if majority of runs use monospace font
                            monospace_chars = 0
                            total_chars = 0
                            for run in para.runs:
                                if run.text.strip():
                                    chars = len(run.text)
                                    total_chars += chars
                                    if run.font.name and any(code_font.lower() in run.font.name.lower() 
                                                            for code_font in code_fonts):
                                        monospace_chars += chars
                            
                            # If >50% monospace, treat as code
                            if total_chars > 0 and monospace_chars / total_chars > 0.5:
                                is_code_block = True
                        
                        if is_code_block:
                            # Add as code block
                            if markdown_lines and markdown_lines[-1] != "":
                                markdown_lines.append("")
                            markdown_lines.append("```")
                            markdown_lines.append(text)
                            markdown_lines.append("```")
                            markdown_lines.append("")
                        else:
                            # Check for inline formatting (bold/italic)
                            formatted_text = ""
                            
                            # IMPROVED: Process runs individually to preserve mixed formatting
                            if para.runs:
                                for run in para.runs:
                                    run_text = run.text
                                    if not run_text:
                                        continue
                                    
                                    # Check all formatting for this run
                                    formatted = run_text
                                    
                                    # Underline
                                    if run.font.underline:
                                        formatted = f"<u>{formatted}</u>"
                                    
                                    # Strikethrough
                                    if run.font.strike:
                                        formatted = f"~~{formatted}~~"
                                    
                                    # Superscript
                                    if run.font.superscript:
                                        formatted = f"<sup>{formatted}</sup>"
                                    
                                    # Subscript
                                    if run.font.subscript:
                                        formatted = f"<sub>{formatted}</sub>"
                                    
                                    # Bold and Italic
                                    if run.bold and run.italic:
                                        formatted = f"***{formatted}***"
                                    elif run.bold:
                                        formatted = f"**{formatted}**"
                                    elif run.italic:
                                        formatted = f"*{formatted}*"
                                    
                                    # Text color (convert to HTML span with color)
                                    if run.font.color and run.font.color.rgb:
                                        rgb = run.font.color.rgb
                                        hex_color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
                                        formatted = f'<span style="color: {hex_color}">{formatted}</span>'
                                    
                                    formatted_text += formatted
                            else:
                                formatted_text = text
                            
                            markdown_lines.append(formatted_text)
                            markdown_lines.append("")  # Space after paragraph
                
                elif isinstance(element, CT_Tbl):
                    # Table
                    table = Table(element, doc)
                    
                    # Add spacing before table
                    if markdown_lines and markdown_lines[-1] != "":
                        markdown_lines.append("")
                    
                    for row_idx, row in enumerate(table.rows):
                        cells = []
                        for cell in row.cells:
                            # Clean cell text - remove extra whitespace
                            cell_text = ' '.join(cell.text.split())
                            cells.append(cell_text)
                        
                        # Create markdown table row
                        markdown_lines.append("| " + " | ".join(cells) + " |")
                        
                        # Add separator after header row
                        if row_idx == 0:
                            separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                            markdown_lines.append(separator)
                    
                    # Add spacing after table
                    markdown_lines.append("")
            
            # Clean up multiple empty lines
            final_lines = []
            prev_empty = False
            for line in markdown_lines:
                if line == "":
                    if not prev_empty:
                        final_lines.append(line)
                    prev_empty = True
                else:
                    final_lines.append(line)
                    prev_empty = False
            
            # Write to file with UTF-8-sig encoding
            with open(output_file, 'w', encoding='utf-8-sig') as f:
                f.write('\n'.join(final_lines))
            
            logger.info(f"Successfully converted DOCX to Markdown: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'docx',
                'markdown',
                warnings=warnings
            )
            
        except Exception as e:
            logger.error(f"DOCX to Markdown conversion failed: {e}")
            raise
    
    def _docx_to_html(self, input_file: str, output_file: str, **options) -> ConversionResult:
        """Convert DOCX to HTML"""
        logger.info(f"Converting DOCX to HTML: {input_file} -> {output_file}")
        
        try:
            html_content = self._extract_docx_as_html(input_file)
            
            # Write to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"Successfully converted DOCX to HTML: {output_file}")
            return self._create_success_result(
                input_file,
                output_file,
                'docx',
                'html'
            )
            
        except Exception as e:
            logger.error(f"DOCX to HTML conversion failed: {e}")
            raise
    
    def _extract_docx_as_html(self, input_file: str) -> str:
        """Extract DOCX content as HTML with enhanced styling"""
        doc = Document(input_file)
        
        html_parts = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            '<meta charset="UTF-8">',
            '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
            '<title>Converted Document</title>',
            '<style>',
            '* { box-sizing: border-box; }',
            'body {',
            '    font-family: "Segoe UI", "Calibri", Arial, sans-serif;',
            '    margin: 0;',
            '    padding: 40px;',
            '    line-height: 1.8;',
            '    color: #333;',
            '    background-color: #f5f5f5;',
            '}',
            '.container {',
            '    max-width: 900px;',
            '    margin: 0 auto;',
            '    background: white;',
            '    padding: 60px;',
            '    box-shadow: 0 2px 10px rgba(0,0,0,0.1);',
            '}',
            'h1, h2, h3, h4, h5, h6 {',
            '    color: #1a1a1a;',
            '    margin-top: 32px;',
            '    margin-bottom: 16px;',
            '    font-weight: 600;',
            '    line-height: 1.3;',
            '}',
            'h1 {',
            '    font-size: 2.5em;',
            '    border-bottom: 3px solid #0066cc;',
            '    padding-bottom: 12px;',
            '    margin-bottom: 24px;',
            '}',
            'h2 {',
            '    font-size: 2em;',
            '    border-bottom: 2px solid #e0e0e0;',
            '    padding-bottom: 8px;',
            '    margin-bottom: 20px;',
            '}',
            'h3 { font-size: 1.5em; }',
            'h4 { font-size: 1.25em; }',
            'h5 { font-size: 1.1em; }',
            'h6 { font-size: 1em; font-weight: 600; }',
            'p {',
            '    margin: 16px 0;',
            '    text-align: justify;',
            '}',
            'ul, ol {',
            '    margin: 16px 0;',
            '    padding-left: 30px;',
            '}',
            'li {',
            '    margin: 8px 0;',
            '    line-height: 1.6;',
            '}',
            'table {',
            '    border-collapse: collapse;',
            '    width: 100%;',
            '    margin: 24px 0;',
            '    background-color: white;',
            '    box-shadow: 0 1px 3px rgba(0,0,0,0.1);',
            '}',
            'table, th, td {',
            '    border: 1px solid #ddd;',
            '}',
            'th, td {',
            '    padding: 14px 18px;',
            '    text-align: left;',
            '}',
            'th {',
            '    background-color: #0066cc;',
            '    color: white;',
            '    font-weight: 600;',
            '    text-transform: uppercase;',
            '    font-size: 0.9em;',
            '    letter-spacing: 0.5px;',
            '}',
            'tr:nth-child(even) { background-color: #f9f9f9; }',
            'tr:hover { background-color: #f0f7ff; }',
            'img {',
            '    max-width: 100%;',
            '    height: auto;',
            '    margin: 20px 0;',
            '    border-radius: 4px;',
            '    box-shadow: 0 2px 6px rgba(0,0,0,0.1);',
            '}',
            'strong { font-weight: 600; color: #1a1a1a; }',
            'em { font-style: italic; color: #444; }',
            'code {',
            '    background-color: #f4f4f4;',
            '    padding: 2px 6px;',
            '    border-radius: 3px;',
            '    font-family: "Courier New", monospace;',
            '    font-size: 0.9em;',
            '    color: #c7254e;',
            '}',
            '.list-item { margin: 8px 0; }',
            '@media print {',
            '    body { background: white; padding: 0; }',
            '    .container { box-shadow: none; padding: 40px; }',
            '}',
            '@media (max-width: 768px) {',
            '    .container { padding: 30px 20px; }',
            '    body { padding: 20px; }',
            '}',
            '</style>',
            '</head>',
            '<body>',
            '<div class="container">',
        ]
        
        # Track list state
        in_list = False
        list_type = None  # 'ul' or 'ol'
        
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text = para.text.strip()
                
                if not text:
                    # Close any open list
                    if in_list:
                        html_parts.append(f'</{list_type}>')
                        in_list = False
                        list_type = None
                    continue
                
                # Detect heading
                if para.style.name.startswith('Heading'):
                    # Close any open list
                    if in_list:
                        html_parts.append(f'</{list_type}>')
                        in_list = False
                        list_type = None
                    
                    level = para.style.name.replace('Heading ', '')
                    try:
                        level = int(level)
                        if level > 6:
                            level = 6
                        html_parts.append(f'<h{level}>{text}</h{level}>')
                    except:
                        html_parts.append(f'<p>{text}</p>')
                
                # IMPROVED: Detect list items with proper wrapping
                elif para.style.name.startswith('List'):
                    is_bullet = 'Bullet' in para.style.name
                    is_number = 'Number' in para.style.name
                    
                    if is_bullet or is_number:
                        current_list_type = 'ul' if is_bullet else 'ol'
                        
                        # Open list if needed or if list type changed
                        if not in_list or list_type != current_list_type:
                            if in_list:
                                html_parts.append(f'</{list_type}>')
                            html_parts.append(f'<{current_list_type}>')
                            in_list = True
                            list_type = current_list_type
                        
                        html_parts.append(f'<li>{text}</li>')
                    else:
                        # Unknown list type, treat as paragraph
                        if in_list:
                            html_parts.append(f'</{list_type}>')
                            in_list = False
                            list_type = None
                        html_parts.append(f'<p>{text}</p>')
                
                else:
                    # Close any open list
                    if in_list:
                        html_parts.append(f'</{list_type}>')
                        in_list = False
                        list_type = None
                    
                    # IMPROVED: Check for code blocks (monospace fonts)
                    is_code = False
                    if para.runs:
                        code_fonts = ['Courier New', 'Consolas', 'Monaco', 'Menlo', 'Courier']
                        monospace_count = sum(1 for run in para.runs 
                                            if run.font.name and 
                                            any(cf.lower() in run.font.name.lower() for cf in code_fonts))
                        if monospace_count > len(para.runs) / 2:
                            is_code = True
                    
                    if is_code:
                        html_parts.append(f'<pre><code>{text}</code></pre>')
                    else:
                        # IMPROVED: Process runs for mixed formatting
                        formatted_html = ""
                        if para.runs:
                            for run in para.runs:
                                run_text = run.text
                                if not run_text:
                                    continue
                                
                                # Start with the base text
                                formatted = run_text
                                
                                # Apply all formatting layers
                                # Underline
                                if run.font.underline:
                                    formatted = f'<u>{formatted}</u>'
                                
                                # Strikethrough
                                if run.font.strike:
                                    formatted = f'<del>{formatted}</del>'
                                
                                # Superscript
                                if run.font.superscript:
                                    formatted = f'<sup>{formatted}</sup>'
                                
                                # Subscript
                                if run.font.subscript:
                                    formatted = f'<sub>{formatted}</sub>'
                                
                                # Bold and Italic
                                if run.bold and run.italic:
                                    formatted = f'<strong><em>{formatted}</em></strong>'
                                elif run.bold:
                                    formatted = f'<strong>{formatted}</strong>'
                                elif run.italic:
                                    formatted = f'<em>{formatted}</em>'
                                
                                # Text color
                                if run.font.color and run.font.color.rgb:
                                    rgb = run.font.color.rgb
                                    hex_color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
                                    formatted = f'<span style="color: {hex_color}">{formatted}</span>'
                                
                                formatted_html += formatted
                        else:
                            formatted_html = text
                        
                        html_parts.append(f'<p>{formatted_html}</p>')
            
            elif isinstance(element, CT_Tbl):
                # Close any open list
                if in_list:
                    html_parts.append(f'</{list_type}>')
                    in_list = False
                    list_type = None
                
                table = Table(element, doc)
                html_parts.append('<table>')
                
                for row_idx, row in enumerate(table.rows):
                    html_parts.append('<tr>')
                    
                    for cell in row.cells:
                        # Clean cell text
                        cell_text = ' '.join(cell.text.split())
                        tag = 'th' if row_idx == 0 else 'td'
                        html_parts.append(f'<{tag}>{cell_text}</{tag}>')
                    
                    html_parts.append('</tr>')
                
                html_parts.append('</table>')
        
        # Close any remaining open list
        if in_list:
            html_parts.append(f'</{list_type}>')
        
        html_parts.extend(['</div>', '</body>', '</html>'])
        
        return '\n'.join(html_parts)
