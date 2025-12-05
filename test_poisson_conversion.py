"""
Test script for 2D_Poisson_FEM.pdf presentation conversion
This script tests various conversion methods and evaluates quality.
"""
import os
import sys
import time
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

from converters import UniversalConverter
from converters.pdf_converter import PDFConverter
from ai.quality_checker import QualityChecker
from utils.logger import logger
import fitz

# Test configuration
PDF_FILE = "2D_Poisson_FEM.pdf"
OUTPUT_DIR = Path("test_outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

def analyze_pdf():
    """Analyze PDF structure and content"""
    print("\n" + "="*60)
    print("📊 PDF ANALİZİ")
    print("="*60)
    
    doc = fitz.open(PDF_FILE)
    
    print(f"📄 Dosya: {PDF_FILE}")
    print(f"📑 Sayfa Sayısı: {len(doc)}")
    print(f"📐 Boyut: {doc[0].rect.width:.0f}x{doc[0].rect.height:.0f} (sunum formatı)")
    
    total_text = 0
    total_images = 0
    has_math = False
    
    for i, page in enumerate(doc):
        text = page.get_text()
        total_text += len(text)
        images = page.get_images()
        total_images += len(images)
        
        # Check for mathematical content
        if any(sym in text for sym in ['∆', '∂', '∇', '∫', '∑', '→', 'Ω', '∈']):
            has_math = True
    
    print(f"📝 Toplam Metin: {total_text:,} karakter")
    print(f"🖼️ Toplam Görsel: {total_images}")
    print(f"🔢 Matematik İçeriği: {'Evet ✅' if has_math else 'Hayır'}")
    
    # Sample content from first 3 pages
    print("\n📋 İçerik Önizleme (ilk 3 sayfa):")
    for i in range(min(3, len(doc))):
        page = doc[i]
        text = page.get_text()[:200].replace('\n', ' ')
        print(f"  Sayfa {i+1}: {text}...")
    
    page_count = len(doc)
    doc.close()
    
    return {
        'pages': page_count,
        'total_chars': total_text,
        'total_images': total_images,
        'has_math': has_math
    }


def test_pdf_to_markdown():
    """Test PDF to Markdown conversion"""
    print("\n" + "="*60)
    print("🔄 TEST: PDF → Markdown")
    print("="*60)
    
    output_file = OUTPUT_DIR / "2D_Poisson_FEM_test.md"
    
    converter = PDFConverter()
    start = time.time()
    result = converter.convert(PDF_FILE, str(output_file))
    elapsed = time.time() - start
    
    print(f"⏱️ Süre: {elapsed:.2f} saniye")
    print(f"✅ Başarılı: {result.success}")
    
    if result.success:
        # Read and analyze output
        with open(output_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        print(f"📄 Çıktı Boyutu: {len(content):,} karakter")
        
        # Count elements
        headings = content.count('\n#')
        tables = content.count('|---|')
        math_symbols = sum(1 for sym in ['∆', '∂', '∇', '∫', '∑', '→', 'Ω'] if sym in content)
        
        print(f"📊 Başlıklar: {headings}")
        print(f"📋 Tablolar: {tables}")
        print(f"🔢 Matematik Sembolleri: {math_symbols}")
        
        # Show sample
        print("\n📝 Çıktı Örneği (ilk 500 karakter):")
        print("-" * 40)
        print(content[:500])
        print("-" * 40)
        
        return content
    else:
        print(f"❌ Hata: {result.error}")
        return None


def test_pdf_to_docx():
    """Test PDF to DOCX conversion"""
    print("\n" + "="*60)
    print("🔄 TEST: PDF → DOCX")
    print("="*60)
    
    output_file = OUTPUT_DIR / "2D_Poisson_FEM_test.docx"
    
    converter = PDFConverter()
    start = time.time()
    result = converter.convert(PDF_FILE, str(output_file))
    elapsed = time.time() - start
    
    print(f"⏱️ Süre: {elapsed:.2f} saniye")
    print(f"✅ Başarılı: {result.success}")
    
    if result.success:
        print(f"📄 Dosya Oluşturuldu: {output_file}")
        file_size = output_file.stat().st_size
        print(f"📊 Dosya Boyutu: {file_size:,} bytes")
        
        # Read DOCX and count elements
        from docx import Document
        doc = Document(str(output_file))
        
        paragraphs = len(doc.paragraphs)
        tables = len(doc.tables)
        
        print(f"📝 Paragraflar: {paragraphs}")
        print(f"📋 Tablolar: {tables}")
        
        return True
    else:
        print(f"❌ Hata: {result.error}")
        return False


def test_pdf_to_html():
    """Test PDF to HTML conversion"""
    print("\n" + "="*60)
    print("🔄 TEST: PDF → HTML")
    print("="*60)
    
    output_file = OUTPUT_DIR / "2D_Poisson_FEM_test.html"
    
    converter = PDFConverter()
    start = time.time()
    result = converter.convert(PDF_FILE, str(output_file))
    elapsed = time.time() - start
    
    print(f"⏱️ Süre: {elapsed:.2f} saniye")
    print(f"✅ Başarılı: {result.success}")
    
    if result.success:
        with open(output_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        print(f"📄 Çıktı Boyutu: {len(content):,} karakter")
        
        # Count HTML elements
        h1_count = content.count('<h1')
        h2_count = content.count('<h2')
        table_count = content.count('<table')
        
        print(f"📊 H1 Başlıklar: {h1_count}")
        print(f"📊 H2 Başlıklar: {h2_count}")
        print(f"📋 Tablolar: {table_count}")
        
        return content
    else:
        print(f"❌ Hata: {result.error}")
        return None


def test_quality_check(original_text: str, converted_text: str):
    """Test AI quality check"""
    print("\n" + "="*60)
    print("🤖 KALİTE ANALİZİ")
    print("="*60)
    
    try:
        checker = QualityChecker(method='heuristic')
        
        # Quick quality metrics
        original_len = len(original_text) if original_text else 0
        converted_len = len(converted_text) if converted_text else 0
        
        if original_len > 0:
            retention_ratio = converted_len / original_len
            print(f"📊 Metin Koruma Oranı: {retention_ratio:.2%}")
        
        # Check for important content preservation
        important_terms = [
            'Poisson', 'FEM', 'Finite Element', 
            'equation', 'boundary', 'domain',
            '∆', 'Ω', '∂'
        ]
        
        preserved = 0
        for term in important_terms:
            if term in (converted_text or ''):
                preserved += 1
        
        preservation_rate = preserved / len(important_terms) * 100
        print(f"🎯 Önemli Terim Koruma: {preservation_rate:.0f}% ({preserved}/{len(important_terms)})")
        
        # Structure preservation
        has_headings = '#' in (converted_text or '') or '<h' in (converted_text or '')
        has_math = any(sym in (converted_text or '') for sym in ['∆', '∂', '∇', '∫'])
        
        print(f"📋 Başlık Yapısı: {'Korundu ✅' if has_headings else 'Kayıp ❌'}")
        print(f"🔢 Matematik: {'Korundu ✅' if has_math else 'Kayıp ❌'}")
        
        # Overall quality score
        score = (
            (min(retention_ratio, 1.5) / 1.5 * 30) +  # Text retention (max 30)
            (preservation_rate * 0.4) +  # Important terms (max 40)
            (15 if has_headings else 0) +  # Structure (15)
            (15 if has_math else 0)  # Math (15)
        )
        
        print(f"\n⭐ TOPLAM KALİTE SKORU: {score:.0f}/100")
        
        if score >= 80:
            print("   Değerlendirme: Mükemmel ✨")
        elif score >= 60:
            print("   Değerlendirme: İyi 👍")
        elif score >= 40:
            print("   Değerlendirme: Orta 🔶")
        else:
            print("   Değerlendirme: Düşük ⚠️")
        
        return score
        
    except Exception as e:
        print(f"❌ Kalite kontrolü hatası: {e}")
        return 0


def run_all_tests():
    """Run all conversion tests"""
    print("\n" + "="*60)
    print("🚀 2D_Poisson_FEM.pdf DÖNÜŞÜM TESTLERİ")
    print("="*60)
    
    # 1. Analyze PDF
    pdf_info = analyze_pdf()
    
    # Get original text for comparison
    doc = fitz.open(PDF_FILE)
    original_text = ""
    for page in doc:
        original_text += page.get_text()
    doc.close()
    
    results = {}
    
    # 2. Test Markdown conversion
    md_content = test_pdf_to_markdown()
    if md_content:
        results['markdown'] = test_quality_check(original_text, md_content)
    
    # 3. Test DOCX conversion
    docx_success = test_pdf_to_docx()
    results['docx'] = docx_success
    
    # 4. Test HTML conversion
    html_content = test_pdf_to_html()
    if html_content:
        results['html'] = test_quality_check(original_text, html_content)
    
    # Summary
    print("\n" + "="*60)
    print("📊 SONUÇ ÖZETİ")
    print("="*60)
    
    print(f"📄 PDF Analizi: {pdf_info['pages']} sayfa, {pdf_info['total_chars']:,} karakter")
    print(f"🔄 Markdown: Skor {results.get('markdown', 'N/A')}/100")
    print(f"🔄 DOCX: {'Başarılı ✅' if results.get('docx') else 'Başarısız ❌'}")
    print(f"🔄 HTML: Skor {results.get('html', 'N/A')}/100")
    
    print(f"\n📁 Çıktı Dosyaları: {OUTPUT_DIR}/")
    for f in OUTPUT_DIR.glob("2D_Poisson_FEM_test.*"):
        print(f"   - {f.name}")
    
    return results


if __name__ == "__main__":
    run_all_tests()
